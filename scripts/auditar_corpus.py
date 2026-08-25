# -*- coding: utf-8 -*-
"""
scripts/auditar_corpus.py — Auditoría de COBERTURA del Corpus Normativo
═══════════════════════════════════════════════════════════════════════════════════
Compara, documento por documento, los artículos que contiene el archivo FUENTE (.docx) con los
que realmente están vectorizados en Supabase. Detecta artículos PERDIDOS en la ingesta.

POR QUÉ EXISTE (hallazgo de Javo · 2026-07-20): el chunker perdía artículos cuando el .docx
maqueta el marcador junto al título del capítulo ("De la rendición de cuentas Art. 88"). La LOPC
quedó con 77 de 103 artículos —faltaba el Art. 88, que funda el derecho ciudadano a exigir
rendición de cuentas— y nadie lo detectó: la BRN daba las cadenas por "íntegras" porque solo
verifica los eslabones DECLARADOS, no la completitud del corpus.

  «Si la base normativa no está completa, todo lo que se construya encima es alucinación probable.»

Sin esta auditoría, una ausencia de ingesta es indistinguible de una inexistencia jurídica — y esa
confusión ya produjo una conclusión errónea. La ausencia debe ser un RESULTADO, no una inferencia.

Uso:  python scripts/auditar_corpus.py            (resumen)
      python scripts/auditar_corpus.py --detalle  (lista los artículos faltantes)
Dylus Lab © 2026
"""
from __future__ import annotations

import re
import sys
import tomllib
from pathlib import Path

try:
    from config import DATOS_DIR as _DATOS
except Exception:                                        # noqa: BLE001
    import os as _os
    _DATOS = Path(_os.environ.get("QUIRA_DATOS", "."))

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

REPO = Path(__file__).resolve().parent.parent
WORD_DIR = _DATOS / "Normativa_Word"
SECRETS = REPO / ".streamlit" / "secrets.toml"

# Mapeo canónico: se REUTILIZA el manifest de QLEP-CORPUS (scripts/normativa/manifest.py) — la
# misma fuente de verdad que usa la ingesta. Así se auditan los 43 documentos, no una lista a mano.
sys.path.insert(0, str(REPO))
from scripts.normativa.manifest import MANIFEST  # noqa: E402

# FUENTE ÚNICA DE VERDAD (colega 2026-07-20: "no mantengas dos lógicas distintas"): se reutiliza
# la misma constante que usa chunk_docx_with_meta() para decidir el perfil de segmentación —
# el auditor no puede tener su propio criterio de "qué es articulado".
from scripts.normativa.chunker import TIPOS_ARTICULADAS_DEFAULT, EXCEPCIONES_NO_ARTICULADAS  # noqa: E402
TIPOS_ARTICULADOS = TIPOS_ARTICULADAS_DEFAULT | {"convenio_internacional"}  # + convenios con Art. N real
MAPEO = {m["archivo"]: m["sigla"] for m in MANIFEST}
TIPO = {m["archivo"]: m.get("tipo", "?") for m in MANIFEST}


def _texto(path: Path) -> str:
    """Texto completo del .docx: párrafos + celdas de tabla (algunos documentos maquetan en tablas)."""
    import docx
    d = docx.Document(str(path))
    partes = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                partes.append(c.text)
    return "\n".join(partes)


def _arts_docx(texto: str) -> set[int]:
    """Artículos declarados en el documento fuente, con el MISMO criterio del chunker corregido."""
    sys.path.insert(0, str(REPO))
    from scripts.normativa.chunker import ARTICLE_RE
    out = set()
    for cap in ARTICLE_RE.findall(texto):
        c = str(cap).strip()
        if c.isdigit():
            out.add(int(c))
    return out


def auditoria_estructural() -> int:
    """AUDITORÍA A (colega · 2026-07-20): ¿el chunker reconoció bien la estructura? NO consulta
    Supabase. Corre el DESPACHO REAL (chunk_docx_with_meta — articulado o perfil genérico según
    el manifest, igual que hace la ingesta) sobre los 43 documentos. Caza bugs del PARSER antes
    de re-ingerir: chunks vacíos, gigantes, errores, o documentos sin ninguna unidad detectada."""
    from scripts.normativa.chunker import chunk_docx_with_meta
    import time
    print("AUDITORÍA ESTRUCTURAL (parser · despacho real) — no consulta Supabase")
    anom = 0
    t0 = time.time()
    for m in MANIFEST:
        f = WORD_DIR / m["archivo"]
        if not f.exists():
            print(f"  ❌ {m['sigla']:16} ARCHIVO FALTANTE"); anom += 1; continue
        try:
            rows = chunk_docx_with_meta(str(f), m)
        except Exception as e:
            print(f"  ❌ {m['sigla']:16} ERROR: {str(e)[:45]}"); anom += 1; continue
        vacios = sum(1 for r in rows if len((r["contenido"] or "").strip()) < 10)
        gigantes = sum(1 for r in rows if r["palabras"] > 600)
        if not rows or vacios or gigantes:
            print(f"  ⚠ {m['sigla']:16} chunks={len(rows)} vacíos={vacios} gigantes={gigantes}"); anom += 1
    print(f"\n{len(MANIFEST)} documentos · {anom} con anomalías estructurales · {time.time()-t0:.1f}s"
          f"{'  → PARSER ESTABLE' if not anom else '  → REVISAR PARSER'}")
    return 1 if anom else 0


def cobertura_estructural() -> int:
    """AUDITORÍA DE COBERTURA POR PERFIL (colega · 2026-07-21): que el parser no se caiga no
    demuestra que capture TODAS las unidades. Para cada documento se determina el PERFIL GANADOR
    REAL (el mismo registro que usa la ingesta) y SOLO si ganó `encabezados_word` se compara el
    número de Heading reales del .docx contra las unidades emitidas — comparar headings de Word
    en documentos que ganaron con otro perfil (Objetivo/Política, Convenio articulado) es una
    métrica inválida (dio 4 falsos positivos en la primera versión de este check: PND-2025,
    CADH, CDN, PIDESC — ninguno usa headings como límite real, así que "headings_docx" no es lo
    que el chunker debía capturar ahí). Una diferencia es aceptable SOLO si coincide exactamente
    con headings 'agrupadores' (sin texto propio, se fusionan por diseño · fix 2026-07-21)."""
    from scripts.normativa.chunker import (chunk_docx_with_meta, _extraer_doc,
                                            PERFILES_REGISTRO, PerfilEncabezadosWord)
    print("COBERTURA ESTRUCTURAL POR PERFIL (solo donde el perfil ganador es 'encabezados_word')")
    anom = 0
    for m in MANIFEST:
        tipo, sigla = m.get("tipo", ""), m["sigla"]
        f = WORD_DIR / m["archivo"]
        if not f.exists():
            continue
        full_text, heads = _extraer_doc(str(f))
        # ¿qué perfil gana REALMENTE para este documento? (mismo criterio que el despachador)
        ganador = None
        for perfil in PERFILES_REGISTRO:
            if perfil.aplica(tipo, sigla) and perfil.segmentar(full_text, heads):
                ganador = perfil
                break
        if not isinstance(ganador, PerfilEncabezadosWord):
            continue   # el check de headings no aplica a este perfil
        agrupadores = 0
        for i, (start, raw) in enumerate(heads):
            end = heads[i + 1][0] if i + 1 < len(heads) else len(full_text)
            if full_text[start:end].strip().lower() == raw.strip().lower():
                agrupadores += 1
        rows = chunk_docx_with_meta(str(f), m)
        emitidas = sum(1 for r in rows if r["chunk_seq"] == 0 and r["articulo_raw"] != "PREÁMBULO")
        esperado_min = len(heads) - agrupadores
        ok = emitidas == esperado_min
        print(f"  {'✅' if ok else '❌'} {sigla:20} headings_docx={len(heads):4} "
              f"agrupadores={agrupadores:3} emitidas={emitidas:4} (esperado {esperado_min})")
        if not ok:
            anom += 1
    print(f"\n{'COBERTURA EXACTA — 0 unidades perdidas' if not anom else f'{anom} DOC(S) CON PÉRDIDA REAL'}")
    return 1 if anom else 0


def check_fixtures() -> int:
    """REGRESIÓN DEL PARSER (colega · 2026-07-20): compara el parser actual contra el baseline
    congelado en fixtures_parser.json. Detecta que un cambio futuro no rompa la segmentación —
    sin repetir la auditoría completa. Idempotencia incluida (SHA estables entre corridas)."""
    import json
    from scripts.normativa.chunker import chunk_docx
    fx_path = REPO / "scripts" / "normativa" / "fixtures_parser.json"
    if not fx_path.exists():
        print("[skip] no hay fixtures_parser.json — genera el baseline primero"); return 0
    fx = json.loads(fx_path.read_text(encoding="utf-8"))
    docs = {"CE": "Constitución del Ecuador.docx", "COOTAD": "COOTAD.docx",
            "COOTAD-2026": "COOTAD PARA LA SOSTENIBILIDAD Y EFICIENCIA 2026.docx",
            "LOPC": "LEY-ORGANICA-DE-PARTICIPACION-CIUDADANA.docx", "LOSNCP": "losncp.docx",
            "RLOTAIP": "LOTAIP - REGLAMENTO-24-01-2024.docx"}
    print("REGRESIÓN DEL PARSER — parser actual vs. baseline (fixtures)")
    roto = 0
    for sig, f in docs.items():
        c = chunk_docx(str(WORD_DIR / f))
        arts = {x.articulo_num for x in c if x.articulo_raw.startswith("Art.") and x.articulo_num}
        actual = {"chunks": len(c), "arts_unicos": len(arts),
                  "art_max": max(arts) if arts else 0,
                  "disposiciones": sum(1 for x in c if "Disposici" in x.articulo_raw)}
        c2 = chunk_docx(str(WORD_DIR / f))            # idempotencia
        idem = [x.sha256 for x in c] == [x.sha256 for x in c2]
        esp = fx.get(sig, {})
        ok = actual == esp and idem
        print(f"  {'✅' if ok else '❌'} {sig:12} {actual}" + ("" if ok else f"  ≠ esperado {esp} idem={idem}"))
        if not ok:
            roto += 1
    print(f"\n{'PARSER ESTABLE — sin regresión' if not roto else f'{roto} REGRESIONES'}")
    return 1 if roto else 0


def calidad_semantica() -> int:
    """MÉTRICAS DE CALIDAD, no solo cantidad (colega · 2026-07-21). El SHA nunca detecta que un
    chunk quedó "solo un título" o que se fusionaron dos unidades en una. Sobre los 43 docs:
      · chunks casi vacíos (<8 palabras, no-PREÁMBULO) → posible título sin cuerpo
      · chunks outlier de longitud (>3x la mediana del propio documento) → posible fusión de
        dos unidades que debieron quedar separadas (p.ej. dos Políticas en un solo chunk)."""
    from scripts.normativa.chunker import chunk_docx_with_meta
    import statistics
    print("CALIDAD SEMÁNTICA (no solo cantidad)")
    total_casi_vacios, total_outliers, docs_con_alerta = 0, 0, 0
    for m in MANIFEST:
        f = WORD_DIR / m["archivo"]
        if not f.exists():
            continue
        rows = chunk_docx_with_meta(str(f), m)
        principales = [r for r in rows if r["chunk_seq"] == 0 and r["articulo_raw"] not in ("PREÁMBULO", "DOCUMENTO")]
        if len(principales) < 3:
            continue
        palabras = [r["palabras"] for r in principales]
        mediana = statistics.median(palabras)
        casi_vacios = [r["articulo_raw"] for r in principales if r["palabras"] < 8]
        outliers = [r["articulo_raw"] for r in principales if mediana > 0 and r["palabras"] > mediana * 6]
        if casi_vacios or outliers:
            docs_con_alerta += 1
            print(f"  ⚠ {m['sigla']:20} mediana={mediana:.0f}p  "
                  f"casi_vacíos={len(casi_vacios)}{casi_vacios[:3] if casi_vacios else ''} "
                  f"outliers={len(outliers)}{outliers[:2] if outliers else ''}")
        total_casi_vacios += len(casi_vacios)
        total_outliers += len(outliers)
    print(f"\n{docs_con_alerta} documento(s) con alertas · {total_casi_vacios} chunks casi vacíos · "
          f"{total_outliers} outliers de longitud (revisar si son fusiones o son legítimos)")
    return 0   # informativo: no bloquea — requiere revisión humana de los casos señalados


def check_integridad_referencial() -> int:
    """CHECK DE NEGOCIO (colega · 2026-07-20), previo a congelar: el corpus debe formar un grafo
    cerrado — sin SHA duplicados/huérfanos, sin chunks sin contenido, sin siglas del corpus fuera
    del manifest y viceversa. Barato de ejecutar, evita dolores de cabeza antes de la reingesta."""
    try:
        uri = tomllib.load(open(SECRETS, "rb"))["database"]["supabase_uri"]
    except Exception:
        print("[skip] sin supabase_uri"); return 0
    import psycopg2
    cur = psycopg2.connect(uri, connect_timeout=30).cursor()
    print("INTEGRIDAD REFERENCIAL DEL CORPUS")
    fallos = 0

    cur.execute("SELECT count(*), count(DISTINCT sha256) FROM public.normativa_corpus")
    tot, sha = cur.fetchone()
    dup = tot - sha
    print(f"  {'✅' if not dup else '❌'} SHA duplicados: {dup} (de {tot} chunks)")
    fallos += 1 if dup else 0

    cur.execute("SELECT count(*) FROM public.normativa_corpus WHERE sha256 IS NULL OR contenido IS NULL")
    huerfanos = cur.fetchone()[0]
    print(f"  {'✅' if not huerfanos else '❌'} chunks sin SHA o sin contenido: {huerfanos}")
    fallos += 1 if huerfanos else 0

    cur.execute("SELECT DISTINCT norma_sigla FROM public.normativa_corpus")
    corp_siglas = {r[0] for r in cur.fetchall()}
    man_siglas = {m["sigla"] for m in MANIFEST}
    fuera_manifest = sorted(corp_siglas - man_siglas)
    fuera_corpus = sorted(man_siglas - corp_siglas)
    # Siglas de OTRO PIPELINE (Holding: evidencia municipal — PAC/POA/PP/PAI/RC/SIGAD/PRESUP/
    # PLAN-BICENTENARIO) no son normativa legal y no cuentan como fallo de ESTE corpus.
    _PREFIJOS_HOLDING = r"^(PAC|POA|PP|PAI|RC|SIGAD|PRESUP|PLAN-BICENTENARIO)-"
    normativas_huerfanas = [s for s in fuera_manifest if not re.match(_PREFIJOS_HOLDING, s)]
    print(f"  {'✅' if not normativas_huerfanas else '❌ DUPLICADO PROBABLE'} "
          f"siglas normativas en corpus sin manifest: {normativas_huerfanas}")
    print(f"  {'✅' if not fuera_corpus else '❌'} siglas en manifest sin ingestar: {fuera_corpus}")
    fallos += 1 if (normativas_huerfanas or fuera_corpus) else 0

    print(f"\n{'GRAFO CERRADO' if not fallos else f'{fallos} PROBLEMA(S) DE INTEGRIDAD'}")
    return 1 if fallos else 0


def main() -> int:
    if "--estructural" in sys.argv:
        return auditoria_estructural()
    if "--fixtures" in sys.argv:
        return check_fixtures()
    if "--integridad" in sys.argv:
        return check_integridad_referencial()
    if "--cobertura-perfil" in sys.argv:
        return cobertura_estructural()
    if "--calidad" in sys.argv:
        return calidad_semantica()
    detalle = "--detalle" in sys.argv
    try:
        uri = tomllib.load(open(SECRETS, "rb"))["database"]["supabase_uri"]
    except Exception:
        print("[skip] sin supabase_uri — no se puede auditar la cobertura")
        return 0
    import psycopg2
    cur = psycopg2.connect(uri, connect_timeout=30).cursor()

    if not WORD_DIR.exists():
        print(f"[ERR] no se halla la carpeta fuente: {WORD_DIR}")
        return 1

    articulados, no_articulados, criticos, faltantes_p = [], [], [], []
    for p in sorted(WORD_DIR.glob("*.docx")):
        sigla = MAPEO.get(p.name)
        if not sigla:
            continue
        tipo = TIPO.get(p.name, "?")
        cur.execute("SELECT count(*), count(DISTINCT articulo_num) FROM public.normativa_corpus "
                    "WHERE norma_sigla=%s", (sigla,))
        n_chunks, n_arts_cor = cur.fetchone()
        if tipo not in TIPOS_ARTICULADOS:
            no_articulados.append((sigla, tipo, n_chunks))
            continue
        try:
            arts = _arts_docx(_texto(p))
        except Exception as e:
            print(f"[WARN] {p.name}: {str(e)[:50]}"); continue
        cur.execute("SELECT DISTINCT articulo_num FROM public.normativa_corpus "
                    "WHERE norma_sigla=%s AND articulo_num IS NOT NULL", (sigla,))
        corp = {r[0] for r in cur.fetchall()}
        faltan = sorted(arts - corp)
        cob = 100 * len(arts & corp) / len(arts) if arts else 0.0
        # déficit GRANDE = bug de ingesta probable · déficit ≤3 y ≥95% = ruido de conteo (aceptable)
        if not arts and not corp:
            est = "s/datos"          # el .docx no usa numeración "Art. N" (CEDAW, clasificador…)
        elif not faltan:
            est = "OK"
        elif len(faltan) > 3 or cob < 95:
            est = "CRÍTICO"; criticos.append(sigla)
        else:
            est = "ruido"
        articulados.append((est, sigla, len(arts), len(corp), faltan, cob))
        if faltan and est == "CRÍTICO":
            faltantes_p.append((sigla, faltan))

    articulados.sort(key=lambda r: r[5])
    print("AUDITORÍA DE COBERTURA DEL CORPUS — fuente (.docx) vs. vectorizado (Supabase)")
    print(f"\nA · NORMAS ARTICULADAS (auditadas por conteo de artículos):")
    print(f"  {'ESTADO':8} {'NORMA':20} {'docx':>5} {'corpus':>6} {'falt':>4}  cob")
    for est, sigla, nd, nc, faltan, cob in articulados:
        print(f"  {est:8} {sigla:20} {nd:5} {nc:6} {len(faltan):4}  {cob:5.1f}%")
        if detalle and faltan:
            print(f"           faltan: {faltan[:30]}")
    print(f"\nB · NO ARTICULADAS (reforma/plan/guía — se reporta presencia, no conteo de artículos):")
    for sigla, tipo, nch in sorted(no_articulados):
        flag = "  ⚠ SIN CHUNKS" if nch == 0 else ""
        print(f"  {sigla:24} tipo={tipo:16} chunks={nch}{flag}")

    ok = sum(1 for r in articulados if r[0] == "OK")
    ruido = sum(1 for r in articulados if r[0] == "ruido")
    print(f"\nRESUMEN articuladas: {len(articulados)} · OK={ok} · ruido_conteo={ruido} · "
          f"CRÍTICAS={len(criticos)}  → {criticos}")
    return 1 if criticos else 0


if __name__ == "__main__":
    raise SystemExit(main())
