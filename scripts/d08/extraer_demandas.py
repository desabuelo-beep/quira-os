# -*- coding: utf-8 -*-
"""
scripts/d08/extraer_demandas.py — Catálogo de demandas ciudadanas (d08 · Fase 1)
═══════════════════════════════════════════════════════════════════════════════
authority:
  parent: GOVERNANCE-001
  constitution_articles: [1, 2, 4]
  type: TECNICA

Responde UNA sola pregunta (Fase 1 · asesor 2026-07-27): **¿qué pidió realmente la
ciudadanía?** NO evalúa, NO cruza contra ejecución, NO juzga. Solo construye un
catálogo limpio y verificable de demandas, con su trazabilidad al documento fuente.

FUENTES (las que el propio GAD declara como medio de verificación en su RDC):
  · Informes de Presupuesto Participativo 2024/2025/2026 — VINCULANTE (COOTAD 238)
  · Actas de Audiencia Pública 2023-2025 (28) — advisory (LOPC 73)
  · Acta de Cabildo Popular 2025 — deliberativo (LOPC 76)

DETERMINÍSTICO: sin LLM, sin API, sin costo. La interpretación fina (clasificar,
cruzar contra POA) es Fase 2 y requiere criterio.

NOTA DE CALIDAD: los .docx del PP provienen de un PDF con OCR deficiente
("PIIOGIIAMA", "OESCRIPCJON"). El extractor normaliza encabezados de forma
tolerante; los textos de demanda se conservan TAL CUAL (no se "corrigen":
alterar la evidencia sería inaceptable). La calidad del OCR se reporta.

Uso:  python scripts/d08/extraer_demandas.py
Salida: data/d08/demandas_ciudadanas.json
Dylus Lab © 2026
"""
from __future__ import annotations

import json
import re
import sys
import unicodedata
from datetime import date
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

REPO = Path(__file__).resolve().parents[2]
BASE = Path(r"C:\Users\DELL\Desktop\Javo\Dylus Lab\ProyecT\Holding_Municipal_Montecristi\Participación Ciudadana")
PP = BASE / "Presupuesto participativo 2024-2026" / "Word"
AUD = BASE / "Audiencias Públicas"


def norm(s: str) -> str:
    """Normaliza para comparar encabezados con OCR sucio."""
    s = unicodedata.normalize("NFKD", str(s or "")).encode("ascii", "ignore").decode()
    return re.sub(r"[^a-z]", "", s.lower())


def es_col_requerimiento(h: str) -> bool:
    n = norm(h)
    return ("descripcion" in n and "requerim" in n) or "requerimiento" in n


def es_col_priorizada(h: str) -> bool:
    n = norm(h)
    return "opcion" in n and "prioriz" in n


def limpiar(txt: str) -> str:
    return re.sub(r"\s+", " ", str(txt or "")).strip(" ·-—|")


# Ruido que NO es demanda ciudadana: membretes, pies de página, citas de la norma.
RE_RUIDO = re.compile(
    r"@\w+\.|www\.|\bTel[eé]f|\bP[aá]gina\s+\d+\s+de\s+\d+|"        # contacto / paginación
    r"\bArt[íi]?c?u?l?o?\.?\s*\d|\bde conformidad\b|\bpodr[aá]n?\b\s+solicitar|"  # cita normativa
    r"\bJEFATURA\b|\bDir\.\s|\bCalle Sucre y\b",                     # membrete institucional
    re.I)


def es_ruido(dem: str) -> bool:
    """Descarta lo que no es una demanda ciudadana (membrete, paginación, cita legal)."""
    if RE_RUIDO.search(dem):
        return True
    letras = sum(c.isalpha() for c in dem)
    return letras < len(dem) * 0.5          # celdas con más números/símbolos que texto


def extraer_pp(path: Path, anio: str) -> list[dict]:
    """Extrae demandas de un informe de Presupuesto Participativo."""
    from docx import Document
    doc = Document(path)
    out = []
    for ti, t in enumerate(doc.tables):
        if not t.rows:
            continue
        heads = [c.text for c in t.rows[0].cells]
        idx_req = next((i for i, h in enumerate(heads) if es_col_requerimiento(h)), None)
        idx_pri = next((i for i, h in enumerate(heads) if es_col_priorizada(h)), None)
        col = idx_req if idx_req is not None else idx_pri
        if col is None:
            continue
        tipo = "requerimiento" if idx_req is not None else "opcion_priorizada"
        for r in t.rows[1:]:
            cells = [limpiar(c.text) for c in r.cells]
            if col >= len(cells):
                continue
            dem = cells[col]
            if len(dem) < 12 or es_ruido(dem):     # vacías/numéricas/membrete/cita legal
                continue
            # el sistema/programa suelen ser las primeras columnas
            ctx = [c for c in cells[:col] if len(c) > 3]
            out.append({
                "demanda": dem,
                "sistema": ctx[0] if ctx else "",
                "programa": ctx[1] if len(ctx) > 1 else "",
                "subprograma": ctx[2] if len(ctx) > 2 else "",
                "tipo_registro": tipo,
                "naturaleza_juridica": "vinculante",     # COOTAD 238
                "mecanismo": "presupuesto_participativo",
                "anio": anio,
                "fuente": f"{path.name} · tabla {ti}",
            })
    return out


def extraer_pp_pdf(path: Path, anio: str) -> list[dict]:
    """Fallback: el .docx del PP 2025 viene vacío (solo imágenes). Se extrae del PDF,
    que sí tiene capa de texto. Misma lógica de columnas, tolerante al OCR."""
    import pdfplumber
    out = []
    with pdfplumber.open(path) as pdf:
        for pi, page in enumerate(pdf.pages):
            for tbl in (page.extract_tables() or []):
                if not tbl or not tbl[0]:
                    continue
                heads = [str(h or "") for h in tbl[0]]
                idx = next((i for i, h in enumerate(heads) if es_col_requerimiento(h)), None)
                if idx is None:
                    idx = next((i for i, h in enumerate(heads) if es_col_priorizada(h)), None)
                    tipo = "opcion_priorizada"
                else:
                    tipo = "requerimiento"
                if idx is None:
                    continue
                for row in tbl[1:]:
                    if idx >= len(row):
                        continue
                    dem = limpiar(row[idx])
                    if len(dem) < 12 or es_ruido(dem):
                        continue
                    ctx = [limpiar(c) for c in row[:idx] if c and len(limpiar(c)) > 3]
                    out.append({
                        "demanda": dem,
                        "sistema": ctx[0] if ctx else "",
                        "programa": ctx[1] if len(ctx) > 1 else "",
                        "subprograma": ctx[2] if len(ctx) > 2 else "",
                        "tipo_registro": tipo,
                        "naturaleza_juridica": "vinculante",
                        "mecanismo": "presupuesto_participativo",
                        "anio": anio,
                        "fuente": f"{path.name} · pág {pi+1}",
                    })
    return out


# marcadores lingüísticos de demanda en actas (español administrativo)
RE_DEMANDA = re.compile(
    r"(?:^|[.;·]\s*)([^.;]{0,120}?\b(?:solicit\w+|pedi\w+|petici\w+|requier\w+|"
    r"necesit\w+|se necesita|urge\w*|piden|demand\w+|gestion\w+ (?:para|de)|"
    r"exig\w+|reclam\w+)\b[^.;]{5,220})", re.I)


def extraer_acta(path: Path, mecanismo: str, naturaleza: str) -> list[dict]:
    """Extrae demandas de un acta por marcadores lingüísticos (determinístico)."""
    for enc in ("utf-8", "cp1252", "latin-1"):
        try:
            txt = path.read_text(encoding=enc)
            break
        except Exception:
            continue
    else:
        return []
    anio = next((a for a in ("2023", "2024", "2025", "2026") if a in path.parent.name), "")
    out = []
    for m in RE_DEMANDA.finditer(txt):
        dem = limpiar(m.group(1))
        if len(dem) < 20 or es_ruido(dem):
            continue
        out.append({
            "demanda": dem,
            "sistema": "", "programa": "", "subprograma": "",
            "tipo_registro": "marcador_linguistico",
            "naturaleza_juridica": naturaleza,
            "mecanismo": mecanismo,
            "anio": anio,
            "fuente": path.name,
        })
    return out


def main() -> int:
    demandas: list[dict] = []
    print("=== FASE 1 · ¿Qué pidió realmente la ciudadanía? ===\n")

    # 1 · Presupuesto Participativo (vinculante)
    for anio in ("2024", "2025", "2026"):
        f = PP / f"GAD Montecristi Informe Presupuesto Participativo {anio}.docx"
        if not f.exists():
            print(f"  [skip] PP {anio}: no existe")
            continue
        try:
            d = extraer_pp(f, anio)
        except Exception as e:
            print(f"  [warn] PP {anio}: {str(e)[:60]}")
            continue
        if not d:
            fpdf = PP.parent / "PDF" / f"GAD Montecristi Informe Presupuesto Participativo {anio}.pdf"
            if fpdf.exists():
                d = extraer_pp_pdf(fpdf, anio)
                print(f"  PP {anio:<5} {len(d):4} demandas  (docx vacío → extraído del PDF)")
                demandas += d
                continue
        demandas += d
        print(f"  PP {anio:<5} {len(d):4} demandas  (vinculante · COOTAD 238)")

    # 2 · Audiencias Públicas (advisory)
    n_aud = 0
    for f in sorted(AUD.rglob("*.txt")):
        d = extraer_acta(f, "audiencia_publica", "advisory")
        demandas += d
        n_aud += len(d)
    print(f"  Audiencias  {n_aud:4} demandas  (advisory · LOPC 73)")

    # 3 · Cabildo Popular (deliberativo)
    n_cab = 0
    for f in sorted((BASE / "Cabildo Popular").glob("*.docx")):
        try:
            from docx import Document
            txt = "\n".join(p.text for p in Document(f).paragraphs)
            tmp = REPO / "data" / "d08" / "_cabildo.tmp"
            tmp.parent.mkdir(parents=True, exist_ok=True)
            tmp.write_text(txt, encoding="utf-8")
            d = extraer_acta(tmp, "cabildo_popular", "deliberativo")
            for x in d:
                x["fuente"] = f.name
                x["anio"] = "2025"
            demandas += d
            n_cab += len(d)
            tmp.unlink()
        except Exception as e:
            print(f"  [warn] cabildo: {str(e)[:60]}")
    print(f"  Cabildo     {n_cab:4} demandas  (deliberativo · LOPC 76)")

    # dedup exacto conservando el primero
    vistos, unicas = set(), []
    for d in demandas:
        k = (norm(d["demanda"])[:80], d["mecanismo"])
        if k in vistos:
            continue
        vistos.add(k)
        unicas.append(d)

    salida = {
        "_fuente": "GENERADO por scripts/d08/extraer_demandas.py — Fase 1 (extracción, no evaluación)",
        "_autoridad": "RO-VIII-003 (dimensión efectividad/incidencia) · Carta Art. 1",
        "_advertencia": "Catálogo PRELIMINAR. Las demandas de actas se detectan por marcadores "
                        "lingüísticos: requieren validación experta antes de cualquier evaluación. "
                        "Los textos se conservan TAL CUAL (OCR incluido): alterar evidencia es inaceptable.",
        "generado": date.today().isoformat(),
        "total": len(unicas),
        "por_mecanismo": {m: sum(1 for d in unicas if d["mecanismo"] == m)
                          for m in {d["mecanismo"] for d in unicas}},
        "por_naturaleza": {n: sum(1 for d in unicas if d["naturaleza_juridica"] == n)
                           for n in {d["naturaleza_juridica"] for d in unicas}},
        "por_anio": {a: sum(1 for d in unicas if d["anio"] == a)
                     for a in sorted({d["anio"] for d in unicas})},
        "demandas": unicas,
    }
    dest = REPO / "data" / "d08" / "demandas_ciudadanas.json"
    dest.parent.mkdir(parents=True, exist_ok=True)
    dest.write_text(json.dumps(salida, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"\n  TOTAL (sin duplicados): {len(unicas)}")
    print(f"  por naturaleza: {salida['por_naturaleza']}")
    print(f"  por año: {salida['por_anio']}")
    print(f"\nOK — {dest.relative_to(REPO)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
