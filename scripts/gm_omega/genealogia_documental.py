# -*- coding: utf-8 -*-
"""
scripts/gm_omega/genealogia_documental.py — GM-Ω · EXPEDIENTE GENEALÓGICO

    Barre los documentos históricos de `tesis historicas/documentos antiguos/`
    y extrae la cadena metodológica que el corpus posterior NO conserva.

    POR QUÉ. Javo aportó una carpeta con material de abril, y el primer
    documento abierto —`Metodologia_SIAP_ICPI.docx`, TERRA/QUADRUM— ya demostró
    que **abril contiene información que no está representada en el corpus
    posterior**: la regla de `E_i` que esta auditoría había declarado
    `NOT_DETERMINABLE`.

    ⚠️ Y LA REGLA DE ESTE BARRIDO: reconstruir la cadena ENTERA antes de
    reescribir ningún diagnóstico. Corregir `007-B0` con el primer documento y
    volver a corregirlo con el segundo sería hacer el trabajo dos veces — y
    dejar en el expediente dos versiones de la misma historia.

    NO decide cuál versión metodológica es «la correcta». Registra que existen,
    con su fecha y su documento. Eso lo dictamina `011`.

Uso:  python scripts/gm_omega/genealogia_documental.py
Dylus Lab © 2026
"""
from __future__ import annotations

import re
import sys
from datetime import datetime
from pathlib import Path

_RAIZ = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(_RAIZ))

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

_SALIDA = _RAIZ / "docs" / "architecture" / "GM-OMEGA_GENEALOGIA_DOCUMENTAL.md"
_FUENTE = _RAIZ.parent / "tesis historicas" / "documentos antiguos"
# ⚠️ Segunda carpeta, hallada el 2026-09-04 siguiendo una pista del chat
# «Profundo»: contiene el ANEXO L de ABRIL y el anexo 0 de FEBRERO — el
# material más antiguo conservado, y el que reconstruyó la genealogía de E_i.
_FUENTE2 = _RAIZ.parent / "ProyecT" / "Terra archivo historico"

# Los ocho huecos que la auditoría dejó abiertos y que abril podría cerrar.
_TEMAS = {
    "E_i · regla": r"(E[_ᵢi]?\s*[—:-].{0,80}(autonom|fricci)|"
                   r"aut[oó]nomo.{0,40}compartido.{0,40}difuso|"
                   r"E_?i\s*=\s*0[.,](90|75|5))",
    "C_i · origen": r"(C[_ᵢi]?\s*[—:-].{0,80}(trazabilidad|org[aá]nic)|"
                    r"Responsabilidad Org[aá]nica Vinculante)",
    "P_i · evolución": r"(P[_ᵢi]?\s*[—:-].{0,80}(peso|presupuest)|"
                       r"PRESUPUESTO_ASIGNADO)",
    "R_i · evolución": r"(R[_ᵢi]?\s*[—:-].{0,80}(relevancia|normativ)|"
                       r"R_?i_?raw|1[.,]725)",
    "AVEP · origen": r"(AVEP|Gesti[oó]n por Mandato|Ruptura Sist[eé]mica|"
                     r"baremo)",
    "universo 25/66": r"(muestra estrat[eé]gica|25 metas|66 metas|"
                      r"universo operacional)",
    "unidad `i`": r"(unidad de an[aá]lisis|cada meta.{0,60}unidad|"
                  r"n\s*=\s*N[uú]mero de metas)",
    "fórmula": r"(P[_ᵢi]?\s*[×x*]\s*R[_ᵢi]?|Σ\s*\(?\s*P|"
               r"ICPI\s*=\s*\[?Σ)",
    "TERRA/QUADRUM": r"(TERRA|QUADRUM|Quadrum Gov)",
}


def _texto(p: Path) -> str:
    """Texto plano de un .docx —párrafos y tablas— o de un .csv."""
    if p.suffix.lower() == ".csv":
        return p.read_text(encoding="utf-8", errors="replace")
    if p.suffix.lower() != ".docx":
        return ""
    try:
        from docx import Document
        d = Document(str(p))
    except Exception:
        return ""
    partes = [x.text for x in d.paragraphs]
    for t in d.tables:
        for fila in t.rows:
            partes.append(" | ".join(c.text for c in fila.cells))
    return "\n".join(partes)


def barrer() -> list[dict]:
    if not (_FUENTE.exists() or _FUENTE2.exists()):
        return []
    out = []
    fuentes = [q for f in (_FUENTE, _FUENTE2) if f.exists() for q in sorted(f.iterdir())]
    for p in fuentes:
        if p.suffix.lower() not in (".docx", ".csv"):
            continue
        txt = _texto(p)
        if not txt:
            continue
        hallazgos = {}
        for tema, patron in _TEMAS.items():
            enc = []
            for m in re.finditer(patron, txt, re.I):
                ini = max(0, m.start() - 90)
                frag = re.sub(r"\s+", " ", txt[ini:m.end() + 220]).strip()
                if frag not in enc:
                    enc.append(frag)
                if len(enc) >= 3:
                    break
            if enc:
                hallazgos[tema] = enc
        out.append({
            "archivo": p.name,
            "modificado": datetime.fromtimestamp(p.stat().st_mtime).strftime("%Y-%m-%d"),
            "caracteres": len(txt),
            "hallazgos": hallazgos,
        })
    return out


def main() -> int:
    docs = barrer()
    if not docs:
        print("[no determinable] no se encontró la carpeta de documentos "
              f"históricos: {_FUENTE}")
        return 2

    total = sum(len(d["hallazgos"]) for d in docs)
    print(f"documentos barridos: {len(docs)} · temas con material: {total}")
    for d in docs:
        if d["hallazgos"]:
            print(f"  {d['archivo'][:46]:<48} {', '.join(d['hallazgos'])}")

    _escribir(docs)
    print(f"→ {_SALIDA.relative_to(_RAIZ).as_posix()}")
    return 0


def _escribir(docs) -> None:
    o: list[str] = []
    A = o.append

    A("# GM-Ω · EXPEDIENTE GENEALÓGICO DOCUMENTAL")
    A("")
    A("**DERIVADO — no editar a mano.** Lo regenera "
      "`scripts/gm_omega/genealogia_documental.py` barriendo "
      "`tesis historicas/documentos antiguos/`.")
    A("")
    A("> ### Por qué existe")
    A("> Javo aportó la carpeta de documentos antiguos, y el primero que se "
      "abrió —`Metodologia_SIAP_ICPI.docx`, de abril, firmado **Ecosistema "
      "TERRA · Quadrum Gov Tech**— demostró que **abril contiene información "
      "que el corpus posterior no conserva**: la regla de `E_i` que esta "
      "auditoría había declarado `NOT_DETERMINABLE`.")
    A(">")
    A("> **Regla del barrido:** reconstruir la cadena entera **antes** de "
      "reescribir ningún diagnóstico. Corregir `007-B0` con el primer documento "
      "y volver a corregirlo con el segundo dejaría dos versiones de la misma "
      "historia en el expediente.")
    A("")

    A("## ★ El hallazgo que abrió el barrido · `E_i` tiene DOS definiciones")
    A("")
    A("| | Definición **A** | Definición **B** |")
    A("|---|---|---|")
    A("| Documento | `Metodologia_SIAP_ICPI.docx` (abril · TERRA/QUADRUM) | `metodologia.docx` (tesis) |")
    A("| Nombre | **Autonomía Orgánica** | **Fricción de Autonomía** |")
    A("| Qué mide | «el **control del director** sobre la ejecución de la meta» | fricción institucional por **delegación** |")
    A("| Escala | `1.0` autónomo · `0.9` compartido · `0.75` difuso | `1.0` directa · `0.90` convenio · `0.75` adscrita |")
    A("| Base | — | `COOTAD Art. 54` · `NCI 200-04` |")
    A("")
    A("**Misma escala numérica, constructos distintos.** Y el motor declara la "
      "definición **A**, textualmente, en `H12!A4`:")
    A("")
    A("```")
    A("Ei: 1.0=autónomo | 0.9=compartido | 0.75=difuso")
    A("```")
    A("")
    A("### ⚠️ Qué corrige esto, y qué NO")
    A("")
    A("| | |")
    A("|---|---|")
    A("| **CORRIGE** | `007-B0` concluyó que «la regla escrita y los valores "
      "implementados nunca coincidieron». Lo que ocurrió es que **la auditoría "
      "contrastó los valores contra la definición B mientras el motor "
      "implementa la A**. Los cinco casos de entidades adscritas con `E_i ≠ "
      "0,75` no eran incoherencia: eran la respuesta correcta bajo la regla que "
      "el motor sí declara |")
    A("| **NO CORRIGE** | que la aplicación meta a meta siga sin verificarse. "
      "Comprobar que los 25 valores cumplen la definición A exige conocer el "
      "«control del director» por meta, y eso no consta en el libro |")
    A("")
    A("**Estado nuevo de `E_i`** — cuatro capas que no deben colapsarse:")
    A("")
    A("```")
    A("  REGLA                        VERIFICADA      · abril + H12!A4")
    A("  VALORES DEL MOTOR            VERIFICADOS     · 25 literales, estables")
    A("  APLICACIÓN META A META       PENDIENTE       · falta el insumo")
    A("  CORRESPONDENCIA VALOR↔REGLA  POR AUDITAR     · no se puede hoy")
    A("```")
    A("")
    A("⚠️ **La definición B no se descarta.** No era basura: pasa a ser "
      "**evidencia de divergencia o evolución metodológica entre documentos**. "
      "Cuál rige lo dictamina `011`, no este expediente.")
    A("")
    A("Y coincide con la corrección que Javo hizo el 2026-09-03 —*«castigar al "
      "GAD por derivar una obra a EP Aseo no es viable: es la misma "
      "institucionalidad»*—: **la definición A no penaliza el organigrama, mide "
      "el control del director.** Su criterio coincidía con la metodología "
      "original, no con la tesis.")
    A("")

    A("## ★ Y el ICPI original tenía CINCO factores")
    A("")
    A("`QUADRUM_ICPI_Calculadora.csv` conserva la construcción primitiva:")
    A("")
    A("```")
    A("  PRODUCTO     = Pi × Vi × Ei × Ti × Ri      ← CINCO · sin C_i")
    A("  Pi           = PRESUPUESTO_ASIGNADO / PRESUPUESTO_TOTAL")
    A("  Ri           = 0.5 · 1 · 1.5               ← sin normalizar por 1,725")
    A("  DENOMINADOR  = Pi × Ri                     ← igual que hoy")
    A("```")
    A("")
    A("Luego la evolución documentada es:")
    A("")
    A("```")
    A("  ICPI_original = Pi · Vi · Ei · Ti · Ri            (abril · 5 factores)")
    A("  ICPI_actual   = Pi · Ri · Vi · Ei · Ti · Ci       (v5.7 · 6 factores)")
    A("```")
    A("")
    A("**`C_i` no es parte del constructo original: es una incorporación "
      "posterior.** Eso reformula `011-C`, que ya no debe preguntar sólo si "
      "está bien multiplicar seis factores, sino:")
    A("")
    A("> **¿Qué transformación metodológica ocurrió entre el ICPI original y el "
      "actual, qué constructo representa cada versión, y está justificada la "
      "incorporación de `C_i` y la arquitectura algebraica resultante?**")
    A("")
    A("⚠️ Y `R_i` también cambió de representación —de `0.5·1·1.5` crudo a "
      "normalizado por el máximo teórico—. **Eso entra en la genealogía de "
      "`R_i`, no se trata automáticamente como error.**")
    A("")

    A("## ★ La UNIDAD DE ANÁLISIS está declarada — y hay que leerla con cuidado")
    A("")
    A("`ICPI.docx` (la tesis) la declara explícitamente:")
    A("")
    A("> **Unidad de Análisis:** el **flujo informativo y programático** en la "
      "trayectoria `Plan de Trabajo → PDOT → POA → SIGAD` del GAD Municipal de "
      "Montecristi.")
    A("")
    A("⚠️ **Pero eso NO responde todavía a `011-A`, y confundirlo sería el error "
      "de siempre.** Son dos cosas distintas:")
    A("")
    A("| | |")
    A("|---|---|")
    A("| **Unidad de análisis de la INVESTIGACIÓN** | qué fenómeno se estudia → "
      "el flujo/trayectoria. **Declarada** |")
    A("| **Unidad `i` de la FÓRMULA** | qué objeto se indexa en `Σ` y recibe "
      "`P·R·V·E·T·C`. **Sigue sin declararse** |")
    A("")
    A("Que la tesis diga «estudio el flujo» no dice si `i` es una meta, un "
      "agregado o una unidad programática. **Es material valioso para `011-A`, "
      "no su respuesta** — y tratarlo como respuesta sería colapsar el objeto "
      "de estudio con el objeto de cálculo.")
    A("")

    A("## `C_i` · su origen conceptual")
    A("")
    A("`memoriaa algo quira.docx` lo enlaza con la doctrina fundacional:")
    A("")
    A("> **Responsabilidad Orgánica Vinculante (`C_i`)** → «✅ VIVO como… "
      "`H07c`: firma del Director activa `Ti_V`»")
    A("")
    A("Coincide con `TERMINOLOGY_ORIGIN_v1.md`, que lo define como *«la "
      "obligación técnica y legal de que cada meta de inversión esté "
      "ineludiblemente anclada a una unidad administrativa específica […] el "
      "antídoto contra la burocracia diluida»*. `C_i` **sí tiene genealogía "
      "conceptual documentada** — lo que no tiene es presencia en el ICPI "
      "original de cinco factores.")
    A("")

    A("## ⚠️ El barrido de los dos documentos grandes dio NEGATIVO")
    A("")
    A("`historico construccion quira.docx` (358 KB · 6.194 párrafos) y "
      "`memoriaa algo quira.docx` (74 KB) son **historiales de construcción del "
      "software**, no documentos metodológicos: sprints, QLEP, Neo4j, GeoTwin, "
      "TERRA Ciudadana. Se buscaron en ellos:")
    A("")
    A("| Buscado | Resultado |")
    A("|---|---|")
    A("| Primera aparición de la **fórmula de 6 factores** | ❌ nada |")
    A("| Primera aparición de la **fórmula de 5 factores** | ❌ nada |")
    A("| **Definición operacional de `i`** | ❌ nada |")
    A("| Incorporación de `C_i` al producto | ❌ sólo su concepto, ya conocido |")
    A("| Transición TERRA/QUADRUM → QUIRA | 🟡 menciones, sin decisión metodológica |")
    A("")
    A("**Es un resultado, no un fracaso.** Lo que establece es que **el "
      "documento que explica la transición del ICPI de cinco a seis factores no "
      "está en este corpus**. Y la consecuencia es la que el asesor anticipó:")
    A("")
    A("> Si los documentos históricos tampoco resuelven qué representa `i`, "
      "entonces **`011-A` tiene que decidirlo** — no descubrirlo. Y fabricar "
      "una definición retrospectiva para que la fórmula parezca más coherente "
      "de lo que era sería el peor desenlace posible.")
    A("")

    A("## ★★★ LA GENEALOGÍA DE `E_i`, RECONSTRUIDA")
    A("")
    A("El `ANEXO L MANUAL TÉCNICO QUADRUM v5.0` —**3 de abril de 2026**, "
      "hallado en `ProyecT/Terra archivo historico/`— documenta la fórmula y la "
      "función que la implementa:")
    A("")
    A("```")
    A("ICPI = [Σ(Vi × Pi × Ei × Ti × Ri) / Σ(Pi × Ri)] × 100")
    A("")
    A("def calcular_ICPI_dinamico(promesas_df):")
    A("    - Vi: float (0.0, 0.5, 1.0) - Verificación documental")
    A("    - Pi: float               - Peso presupuestario normalizado")
    A("    - Ei: int (1-5)           - ENTIDAD CUSTODIO RESPONSABLE   ⚠️")
    A("    - Ti: float (0.0-1.0)     - Avance temporal")
    A("    - Ri: float (0.5-1.5)     - Relevancia constitucional")
    A("```")
    A("")
    A("> **En abril, `E_i` no era un coeficiente: era un IDENTIFICADOR de "
      "entidad custodio —un entero de 1 a 5— multiplicándose dentro del "
      "producto.**")
    A("")
    A("Y eso es matemáticamente incoherente: una meta ejecutada por la entidad "
      "`5` valdría **cinco veces** más que una idéntica de la entidad `1`, sin "
      "ninguna razón metodológica. El identificador entra en el cálculo como si "
      "fuera una magnitud.")
    A("")
    A("### La transición está documentada · su MOTIVO no")
    A("")
    A("⚠️ **Una versión anterior de este expediente escribió «ahí está por qué "
      "`E_i` cambió: alguien vio que multiplicar por un ID no tenía sentido».** "
      "Eso es una **hipótesis causal**, no un hecho: ninguna fuente dice que se "
      "cambiara por esa razón. Es la explicación elegante que va más allá de la "
      "evidencia — el mismo error que `DOC-009` y `DOC-019` persiguen.")
    A("")
    A("**La formulación correcta:**")
    A("")
    A("> La transición de `E_i` desde identificador de entidad hacia coeficiente "
      "**está documentada**; la **motivación causal** de esa transformación "
      "permanece **NO DETERMINABLE** salvo evidencia explícita.")
    A("")
    A("Y la secuencia observada es ésta:")
    A("")
    A("| # | Estado de `E_i` | Fuente | Fecha |")
    A("|---|---|---|---|")
    A("| 1 | `int (1-5)` · **identificador** de entidad custodio | `ANEXO L QUADRUM v5.0` | **3-abr-2026** |")
    A("| 2 | «Autonomía Orgánica» · `1.0 / 0.9 / 0.75` — **control del director** | `Metodologia_SIAP_ICPI` | abril |")
    A("| 3 | coeficiente `1 · 0.9 · 0.5` en 20 promesas | calculadora QUADRUM | s/f |")
    A("| 4 | «Fricción de Autonomía» · `COOTAD 54 · NCI 200-04` | tesis | s/f |")
    A("| 5 | `1 / 0.90 / 0.75`, citando **«autónomo/compartido/difuso»** | `H12!A4` + 25 literales | v5.7 |")
    A("")
    A("El motor actual cita el estado **2**. La tesis describe el **4**. Y esta "
      "auditoría comparó los valores contra el 4 cuando el motor implementa el "
      "2 — de ahí que «no cuadraran».")
    A("")
    A("⚠️ **Límite de lo afirmado**: el `ANEXO L` **especifica** esa función; "
      "que llegara a ejecutarse con `Ei` entero no está demostrado. Lo "
      "demostrado es qué decía la especificación de abril.")
    A("")

    A("## ★★★ Y LA FÓRMULA ORIGINAL SÍ TENÍA EL `× 100`")
    A("")
    A("```")
    A("abril    ICPI = [Σ(Vi × Pi × Ei × Ti × Ri) / Σ(Pi × Ri)] × 100")
    A("v5.7     H12!B33 = B31/B32                              ← sin ×100")
    A("```")
    A("")
    A("⚠️ **Y aquí también hay que frenar.** Una versión anterior llamó a esto "
      "«la pérdida de un factor». **El cambio de escala está demostrado; su "
      "carácter, no.** Hay dos lecturas y la evidencia no elige entre ellas:")
    A("")
    A("| | |")
    A("|---|---|")
    A("| **A · cambio semántico intencional** | el motor pasó a almacenar el "
      "ICPI como **proporción** (`0,274582`) y la presentación lo convierte. "
      "No hay pérdida matemática: hay cambio de representación interna |")
    A("| **B · pérdida accidental** | `B33` pretendía ser porcentaje y el `×100` "
      "se eliminó sin actualizar superficies ni documentación. Entonces sí es un "
      "defecto de representación |")
    A("")
    A("**Formulación correcta:**")
    A("")
    A("> El `×100` presente en la especificación histórica no está en la "
      "expresión canónica actual de `B33`. **El cambio de escala interna está "
      "demostrado; su carácter intencional o accidental permanece pendiente de "
      "determinación.**")
    A("")
    A("Lo que `007-X` sí probó es que **existe una inconsistencia real de "
      "rotulado** —69 cabeceras imprimen «0,27 %»—, y eso es compatible con "
      "ambas lecturas: en A sería un rótulo mal actualizado; en B, la huella del "
      "factor perdido. La capa API compensa (`H73!ICPI_GLOBAL_PCT = B33*100`) y "
      "por eso la UI publica bien.")
    A("")

    A("## ★★★ `011-A` RESUELTO EN SU GENEALOGÍA · `i` era una PROMESA")
    A("")
    A("El `ANEXO_M_ICPI_DINAMICO_PROFESIONAL` —«Formalización del Algoritmo "
      "ICPI Dinámico», febrero— formaliza:")
    A("")
    A("```")
    A("ICPI(t) = [Σᵢ (Vᵢ(t) × Pᵢ × Eᵢ × Tᵢ(t) × Rᵢ) / Σᵢ (Pᵢ × Rᵢ)] × 100")
    A("")
    A("«Si Vᵢ(t) = 0 → Contribución_PROMESA_i = 0»")
    A("«la regla de anulación documental garantiza que ausencia de evidencia")
    A(" invalida PROMESA independientemente de otros factores»")
    A("```")
    A("")
    A("**`i` indexaba PROMESAS del Plan de Trabajo del CNE**, no metas del PDOT. "
      "Y la calculadora QUADRUM lo confirma desde el dato: su columna se llama "
      "`PROMESA_CNE` y sus identificadores son `A-001`…`A-020`.")
    A("")
    A("### Y el motivo lo declara Javo")
    A("")
    A("> **«Comenzamos tomando el plan CNE como promesa original; luego "
      "replanteamos con PDOT pues era mandato.»**")
    A("> — Javo, 2026-09-04")
    A("")
    A("**No fue deriva: fue una mejora metodológica.** El plan de campaña es una "
      "promesa; el PDOT es el **instrumento vinculante** —`COPFP Art. 41`: «los "
      "planes de desarrollo son las directrices principales»—. Medir la "
      "congruencia contra lo que obliga jurídicamente es más defendible que "
      "medirla contra lo que sólo se ofreció.")
    A("")
    A("Y explica por qué la tesis define la unidad de análisis como la "
      "**trayectoria** `Plan de Trabajo → PDOT → POA → SIGAD`: **la promesa no "
      "desapareció, se convirtió en el eslabón anterior.**")
    A("")
    A("### Estado de `011-A`")
    A("")
    A("| | |")
    A("|---|---|")
    A("| Unidad histórica de `i` | **DEMOSTRADA** · promesa CNE (`ANEXO M` + calculadora) |")
    A("| Motivo del cambio | **DECLARADO** por Javo · el PDOT es mandato |")
    A("| Unidad vigente de `i` | meta del PDOT — **operando, pero sin declaración formal en el canon** |")
    A("| Si una unidad agrega varias metas (`008-R`) | **PENDIENTE** · `011-B` |")
    A("")
    A("⚠️ `011-A` deja de ser una pregunta abierta y pasa a ser **un acto "
      "pendiente**: la unidad ya está decidida y opera; lo que falta es "
      "**declararla en el canon con su genealogía**. Que es un trabajo mucho "
      "menor que decidirla desde cero.")
    A("")

    A("## ★★ POR QUÉ NO HAY DOCUMENTO · la evolución fue conversacional")
    A("")
    A("Javo lo aclaró y resuelve el hueco que 13 documentos no llenaron:")
    A("")
    A("> **«La fórmula del ICPI vino evolucionando desde ENERO, cuando "
      "comenzamos a trabajar con Claude. Antes era sólo yo trabajando, pero al "
      "trabajar con Claude desde el chat —antes de Code— me pudo ayudar a "
      "potenciar la fórmula; por eso ha venido cambiando, es decir, "
      "evolucionado.»**")
    A("> — Javo, 2026-09-04")
    A("")
    A("**No falta un documento: nunca hubo uno.** La entrada de `C_i`, el paso "
      "de cinco a seis factores y la renormalización de `P_i` y `R_i` **no "
      "están documentados porque ocurrieron en diálogo**, iterando sobre la "
      "tesis. El barrido no fracasó — buscaba en el sitio equivocado.")
    A("")
    A("Y la primera iteración tiene nombre: el chat **«profundo»**, donde se "
      "sentaron las bases del ecosistema que hoy es QUIRA, partiendo de la "
      "tesis.")
    A("")
    A("### Lo que esto corrige de esta auditoría")
    A("")
    A("| Se venía diciendo | Lo que es |")
    A("|---|---|")
    A("| «el hueco es abril → mayo» | **el hueco empieza en ENERO** — abril es "
      "sólo donde aparece el primer artefacto conservado |")
    A("| «falta el documento que explica `C_i`» | **no existe tal documento**; "
      "la justificación vive en una conversación |")
    A("| «la divergencia nació con la implementación» (`007-B0`) | nació de una "
      "**evolución iterativa** que el canon no registró |")
    A("")
    A("### Y la deuda que sí queda, nombrada · `DOC-022`")
    A("")
    A("> **Una decisión que sostiene el motor y vive sólo en una conversación es "
      "una decisión fuera del canon.**")
    A("")
    A("No es un reproche al método: iterar con Claude fue lo que potenció la "
      "fórmula, y el motor resultante funciona y está validado empíricamente. "
      "**La deuda es otra**: lo que el sistema ejecuta debe poder explicarse "
      "**desde el sistema**, no desde la memoria de quien lo construyó. Sin eso, "
      "`011` tiene que volver a decidir lo que ya se decidió una vez.")
    A("")
    A("### Qué haría falta del chat «profundo», y para qué exactamente")
    A("")
    A("No hace falta todo: hay **cuatro decisiones** cuya justificación está sin "
      "recuperar, y cada una tiene un destino concreto en la auditoría.")
    A("")
    A("| Buscar | Para | Si no aparece |")
    A("|---|---|---|")
    A("| Por qué entró `C_i` | `011-C` | se decide de nuevo, y se declara decisión nueva |")
    A("| Por qué 5 → 6 factores | `011-C` | ídem |")
    A("| Por qué se renormalizaron `P_i` y `R_i` | `011-C` · `007-A` ya probó que la de `R` es inocua | ídem |")
    A("| Qué es `i` | `011-A` | **`011-A` decide**, no descubre |")
    A("")
    A("⚠️ Y una salvedad de método: lo que aparezca en el chat será **evidencia "
      "de la decisión**, no canon por sí mismo. Para entrar al canon tendrá que "
      "declararse como lo que es —una decisión de diseño, con su fecha y su "
      "motivo— igual que cualquier otra.")
    A("")

    A("## Tabla de genealogía")
    A("")
    A("⚠️ **Las fechas salen del CONTENIDO, no del sistema de ficheros.** Donde "
      "no hay fecha fiable se escribe `NO DETERMINABLE`, nunca el `mtime`.")
    A("")
    A("| Elemento | Versión histórica | Evidencia | Cambio | Justificación | Estado |")
    A("|---|---|---|---|---|---|")
    A("| `P_i` | `PRESUPUESTO_ASIGNADO / TOTAL` | calculadora QUADRUM | → normalizado Σ=1 sobre 25 | **no hallada** | 🟡 cambio documentado, motivo no |")
    A("| `R_i` | `0.5 · 1 · 1.5` crudo | calculadora QUADRUM | → normalizado por máximo 1,725 | **no hallada** | 🟡 ídem |")
    A("| `V_i` | 3 niveles con núcleo financiero | `H13!B16-B21` | regla anterior `suma≥2` → actual | ✅ **documentada en el libro** | 🟢 con límite de reconstrucción |")
    A("| `E_i` | **A**: control del director (abril) · **B**: fricción por delegación (tesis) | `Metodologia_SIAP_ICPI` · `metodologia.docx` · `H12!A4` | dos definiciones coexistentes | **no hallada** | 🟡 regla verificada, aplicación pendiente |")
    A("| `T_i` | ratio por entidad | `H07b` | curva de pacing sustituyó a `mes/12` | **no hallada** · la nota quedó desfasada | 🟡 |")
    A("| `C_i` | **ausente** del ICPI original | calculadora QUADRUM (5 factores) | **incorporación posterior** | ❌ **NO HALLADA** | 🔴 el hueco principal |")
    A("| `i` | «flujo informativo/programático» (unidad de INVESTIGACIÓN) | `ICPI.docx` | — | — | 🔴 unidad de CÁLCULO sin declarar |")
    A("| Fórmula | `Pi·Vi·Ei·Ti·Ri` (5) | calculadora QUADRUM | → `Pi·Ri·Vi·Ei·Ti·Ci` (6) | ❌ **NO HALLADA** | 🔴 |")
    A("| AVEP | eje conceptual → 4 niveles | `TERMINOLOGY_ORIGIN_v1` | → 5 niveles + fórmula `IF` ×11 hojas | 🟡 el incidente sí (`H01!A28`) | 🟡 |")
    A("| Universo | «muestra estratégica» | tesis | → 25 rotuladas `Total_Metas_PDOT` | ✅ criterio: mayor monto (Javo) | 🟢 |")
    A("")
    A("**Fecha interna más antigua localizada en el corpus: `2026-05-16`** "
      "(Gold Master v5.4). El material de abril existe —la metodología "
      "TERRA/QUADRUM— pero **entre abril y el 10 de mayo no hay ningún "
      "artefacto conservado**, y ahí es donde ocurrieron los cambios que esta "
      "tabla no puede justificar.")
    A("")

    A("## Barrido documental")
    A("")
    A("⚠️ **La columna «copiado» NO es la fecha del documento.** Todos los "
      "archivos se trasladaron a esta carpeta el mismo día, así que el sistema "
      "de ficheros dice `2026-09-04` para material que es de enero o de abril. "
      "**La fecha real vive en el contenido** —`Metodologia_SIAP_ICPI.docx` se "
      "identifica como TERRA/QUADRUM © 2026 y su copia en Drive está fechada el "
      "25 de abril—. Ordenar la genealogía por `mtime` produciría una "
      "cronología falsa.")
    A("")
    A("| Documento | Copiado | Caracteres | Temas con material |")
    A("|---|---|---:|---|")
    for d in sorted(docs, key=lambda x: x["modificado"]):
        temas = ", ".join(f"`{t}`" for t in d["hallazgos"]) or "—"
        A(f"| `{d['archivo'][:44]}` | {d['modificado']} | {d['caracteres']} | {temas} |")
    A("")

    for d in sorted(docs, key=lambda x: x["modificado"]):
        if not d["hallazgos"]:
            continue
        A(f"### `{d['archivo']}` · {d['modificado']}")
        A("")
        for tema, frags in d["hallazgos"].items():
            A(f"**{tema}**")
            A("")
            for f in frags:
                A(f"- …{f[:300]}…")
            A("")

    A("## Lo que este expediente NO hace")
    A("")
    A("- **No decide qué versión metodológica rige.** Registra que existen, con "
      "su documento y su fecha. El dictamen es de `011`.")
    A("- **No cierra el frente de `E_i`.** Cambia su estado; la correspondencia "
      "valor↔regla sigue por auditar.")
    A("- **No toca el Gold Master ni recalcula el 27,4582 %.** El baseline sigue "
      "congelado mientras se reconstruye la genealogía.")
    A("- **No descarta la definición B** ni ningún documento histórico: la "
      "divergencia entre versiones **es** el objeto de estudio.")
    A("")
    A("---")
    A(f"*GM-Ω · Expediente genealógico · {len(docs)} documentos barridos · "
      "el Gold Master no se modificó · Dylus Lab © 2026*")

    _SALIDA.write_text("\n".join(o) + "\n", encoding="utf-8")


if __name__ == "__main__":
    raise SystemExit(main())
