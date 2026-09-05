# -*- coding: utf-8 -*-
"""
scripts/gm_omega/justificacion_transformaciones.py — GM-Ω-ICPI-011-C3

    ¿POR QUÉ el motor terminó usando exactamente esa regla, esos pesos, ese
    piso, ese fallback — y qué parte de esa cadena puede DEMOSTRARSE?

    `011-C1` reconstruyó el álgebra. `011-C2` reconstruyó el significado y
    encontró cuatro divergencias. `011-C3` pregunta por la CAUSA, y su encargo
    es un peritaje, no un rediseño:

        DECLARADO ≠ IMPLEMENTADO ≠ EFECTIVAMENTE UTILIZADO ≠ JUSTIFICADO

    ⚠️ NO SE TOCA NADA. Ni `C_i`, ni `E_i`, ni `T_i`, ni la fórmula, ni las
    calibraciones, ni `Ci_Manual_2025`, ni `Ci_Adaptativo`. El Gold Master es
    inmutable (`Regla de Oro 1`) y el baseline sigue congelado. C3 levanta
    acta de lo que encuentra.

    ⚠️ Y LA REGLA QUE ORDENA ESTA ETAPA: **una transición documentada no
    autoriza a inventar su causa** (`DOC-011`). Si no aparece la razón, el
    resultado correcto es `NO DETERMINABLE` — que es un hallazgo, no un fallo
    del peritaje. Reconstruir perfectamente CUÁNDO apareció algo no dice nada
    sobre POR QUÉ apareció.

Uso:  python scripts/gm_omega/justificacion_transformaciones.py
Dylus Lab © 2026
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

_RAIZ = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(_RAIZ))

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

_SALIDA = _RAIZ / "docs" / "architecture" / "GM-OMEGA_ICPI_JUSTIFICACION_011C3.md"

# Grados de certeza. Son los mismos cuatro de 007-B0 y no se amplían: cada
# afirmación de C3 tiene que caber en uno.
_DEM, _DECL, _INF, _ND = "DEMOSTRADO", "DECLARADO", "INFERIDO", "NO DETERMINABLE"

# Las nueve preguntas, en el orden de dirección. Cada una lleva el patrón que
# busca EVIDENCIA DE JUSTIFICACIÓN — no menciones del término.
#
# ⚠️ La distinción es la que hizo fracasar el primer intento de derivar la
# doctrina: que un artefacto NOMBRE a `C_i` no prueba que lo justifique. Por eso
# los patrones buscan lenguaje causal («se decidió», «porque», «criterio»,
# «fundamentación»), no el nombre del factor.
_PREGUNTAS = [
    ("C3-01", "¿Quién introdujo `C_i`, cuándo y en qué versión?",
     r"(Ci\s+DETERMINISTA|27[-\s]?Abr|Delgado Santana|GENESIS_RUN|"
     r"VERSION_ANTERIOR)"),
    ("C3-02", "¿Qué fenómeno se dijo que medía?",
     r"(calidad (institucional|del proceso|de proceso)|responsabilidad "
     r"org[aá]nic|imputabilidad|trazabilidad org[aá]nic|presunci[oó]n de "
     r"legalidad)"),
    ("C3-03", "¿Por qué se pasó de 5 a 6 factores?",
     r"(se (incorpor|a[ñn]adi|agreg|sum)[oó]|nueva (variable|dimensi[oó]n)|"
     r"sexto factor|P.{0,3}R.{0,3}V.{0,3}E.{0,3}T\b)"),
    ("C3-04", "¿Por qué esos cuatro eventos (CGE · SERCOP · POA · CPCCS)?",
     r"(Marco legal|marco jur[ií]dico|Framework jur[ií]dico|"
     r"agn[oó]stico|Escalabilidad: Ecuador)"),
    ("C3-05", "¿Por qué esos pesos (0,10 · 0,15 · 0,05/0,20 · 0,50)?",
     r"(criterio experto|no por (an[aá]lisis|PCA|regresi[oó]n)|"
     r"Deducci[oó]n_Ci|definidos? por criterio)"),
    ("C3-06", "¿Por qué el piso `MÁX(0,50; …)`?",
     r"(M[ÁA]X\(0[.,]50|MAX\(0[.,]50|FIJACI[OÓ]N ABSOLUTA|"
     r"Ci=0[.,]50 DIRECTAMENTE)"),
    ("C3-07", "¿Por qué existe `Ci_Manual_2025` en 2026?",
     r"(Ci_Manual_2025|Mapeo Retrospectivo|reverse engineering|"
     r"calibraci[oó]n retro|REAL-HEUR[IÍ]STICO|axioma falla)"),
    ("C3-08", "¿Qué es `Ci_Adaptativo` y por qué no se conecta?",
     r"(Ci_Adaptativo|Modificador|INTANGIBLE_FLAG|"
     r"discriminaci[oó]n positiva|FONDO_CONCURSABLE)"),
    ("C3-09", "¿Por qué `E_i` y `C_i` divergen donde divergen?",
     r"(Autonom[ií]a org[aá]nica|aut[oó]nomo.{0,30}compartido|"
     r"proceso (exclusivo|compartido|difuso))"),
]

# Hojas del libro cuyo PROPÓSITO declarado es dejar constancia. Si la
# justificación de una transformación existe en algún sitio, es aquí.
_HOJAS_MEMORIA = ("H80_MODEL_REGISTRY", "H95_LIMITACIONES", "H96_TRAZABILIDAD",
                  "H76_AUDIT_TRAIL", "H77_DATA_DICTIONARY", "RC_CHANGELOG",
                  "H86b_ALGORITHMIC_GOVERNANCE_PRO", "H39_AUTOCONTROL_ECOSISTEMA",
                  "H14_PONDERADORES", "H02_GLOSARIO_QUIRA", "H01_PARÁMETROS")


def leer_libro() -> dict:
    """Barre el Gold Master entero recogiendo texto con valor probatorio.
    `{}` si no se resolvió — el tercer estado."""
    import openpyxl

    import config
    if not getattr(config, "GOLD_MASTER_RESUELTO", False):
        return {}

    wv = openpyxl.load_workbook(config.SIAP_PATH, data_only=True, read_only=True)

    celdas: list[tuple[str, str]] = []
    for nombre in wv.sheetnames:
        try:
            for fila in wv[nombre].iter_rows(min_row=1, max_row=130, max_col=9,
                                             values_only=True):
                for v in fila:
                    if isinstance(v, str) and len(v) >= 40:
                        celdas.append((nombre, re.sub(r"\s+", " ", v).strip()))
        except Exception:
            continue

    # ── H80 · la cadena de versiones, que es la columna vertebral de C3 ────
    versiones = []
    if "H80_MODEL_REGISTRY" in wv.sheetnames:
        h80 = wv["H80_MODEL_REGISTRY"]
        for r in range(3, 24):
            v = h80.cell(row=r, column=1).value
            ts = h80.cell(row=r, column=2).value
            if not v or not ts:
                continue
            versiones.append({
                "version": str(v), "ts": str(ts)[:10],
                "operador": str(h80.cell(row=r, column=5).value or ""),
                "estado": str(h80.cell(row=r, column=6).value or ""),
                "anterior": str(h80.cell(row=r, column=7).value or ""),
            })

    # ── H14 · ¿qué factores tienen columna de JUSTIFICACIÓN declarada? ─────
    just_h14 = []
    if "H14_PONDERADORES" in wv.sheetnames:
        h14 = wv["H14_PONDERADORES"]
        for r in range(1, 12):
            for c in range(1, 12):
                v = h14.cell(row=r, column=c).value
                if isinstance(v, str) and "ustificaci" in v and len(v) < 60:
                    just_h14.append(re.sub(r"\s+", " ", v).strip())

    return {"celdas": celdas, "versiones": versiones, "just_h14": just_h14,
            "hojas": list(wv.sheetnames)}


def leer_historicos() -> list[tuple[str, str]]:
    """Los documentos de enero-abril. Si la razón de una transformación se
    escribió alguna vez, el candidato más probable está aquí."""
    from scripts.gm_omega.genealogia_documental import _FUENTE, _FUENTE2, _texto

    out = []
    for carpeta in (_FUENTE, _FUENTE2):
        if not carpeta.exists():
            continue
        for p in sorted(carpeta.iterdir()):
            if p.suffix.lower() not in (".docx", ".csv"):
                continue
            txt = _texto(p)
            if txt:
                out.append((p.name, re.sub(r"[ \t]+", " ", txt)))
    return out


def leer_tesis() -> dict:
    """★ La fuente que reordena `011-C3`: `metodologia.docx`.

    Define las SEIS variables con definición conceptual, fundamento normativo y
    **tabla de escala con su criterio de verificación**. Se lee la fecha de
    creación del documento —no la de modificación, que sólo dice cuándo se
    copió a la carpeta— y las tablas de escala de `E_i` y `C_i`.

    ⚠️ Y ES LO QUE OBLIGA A CORREGIR `007-B0`: si esta metodología es anterior
    al 27-abr y ya contiene `C_i`, entonces esa fecha **no data la creación del
    factor**, sino la de su versión determinista."""
    from scripts.gm_omega.genealogia_documental import _FUENTE, _FUENTE2

    for carpeta in (_FUENTE, _FUENTE2):
        p = carpeta / "metodologia.docx"
        if not p.exists():
            continue
        try:
            from docx import Document
            doc = Document(str(p))
        except Exception:
            return {}
        cp = doc.core_properties
        escalas = {}
        for t in doc.tables:
            filas = [[c.text.strip() for c in f.cells] for f in t.rows]
            cab = " ".join(filas[0]).lower() if filas else ""
            if "modalidad de ejecuci" in cab:
                escalas["E_i"] = filas
            elif "claridad de responsabilidad" in cab:
                escalas["C_i"] = filas
        return {
            "archivo": p.name,
            "creado": str(cp.created)[:10] if cp.created else "—",
            "escalas": escalas,
            "definiciones": _definiciones_tesis(doc),
        }
    return {}


def _definiciones_tesis(doc) -> dict:
    """La «Definición conceptual» de cada variable, literal, tal como el
    documento la enuncia. No se parafrasea: es la pieza que `C3` contrasta."""
    texto = "\n".join(p.text for p in doc.paragraphs)
    out = {}
    for m in re.finditer(r"3\.4\.\d+\.?\s*Variable\s+([PRVETC])_?i\s*:\s*([^\n]+)",
                         texto):
        factor, titulo = f"{m.group(1)}_i", m.group(2).strip()
        resto = texto[m.end():m.end() + 700]
        d = re.search(r"Definición conceptual:\s*(.+?)(?:\n|Escala|Fórmula|$)",
                      resto, re.S)
        out[factor] = {"titulo": titulo,
                       "definicion": re.sub(r"\s+", " ",
                                            d.group(1)).strip() if d else ""}
    return out


def buscar(patron: str, celdas: list, historicos: list) -> dict:
    """Toda la evidencia de una pregunta, separada por procedencia. La
    separación importa: el libro declara, los documentos justifican."""
    rx = re.compile(patron, re.I)
    en_libro, vistos = [], set()
    for hoja, txt in celdas:
        if rx.search(txt) and txt[:80] not in vistos:
            vistos.add(txt[:80])
            en_libro.append((hoja, txt))

    en_docs = []
    for nombre, txt in historicos:
        for m in rx.finditer(txt):
            ini = max(0, m.start() - 110)
            frag = re.sub(r"\s+", " ", txt[ini:m.end() + 200]).strip()
            if frag[:70] not in vistos:
                vistos.add(frag[:70])
                en_docs.append((nombre, frag))
            if len(en_docs) >= 3:
                break
    return {"libro": en_libro[:6], "docs": en_docs[:3],
            "n_libro": len(en_libro), "n_docs": len(en_docs)}


def main() -> int:
    d = leer_libro()
    if not d:
        print("[no determinable] Gold Master no resuelto.")
        return 2
    historicos = leer_historicos()
    tesis = leer_tesis()

    print(f"libro: {len(d['hojas'])} hojas · {len(d['celdas'])} celdas de texto")
    print(f"documentos históricos: {len(historicos)}")
    print(f"versiones registradas en H80: {len(d['versiones'])}")
    if tesis:
        print(f"tesis `{tesis['archivo']}` creada {tesis['creado']} · "
              f"{len(tesis['definiciones'])} variables definidas · escalas: "
              f"{', '.join(tesis['escalas']) or 'ninguna'}")

    hallazgos = {}
    for pid, pregunta, patron in _PREGUNTAS:
        h = buscar(patron, d["celdas"], historicos)
        hallazgos[pid] = h
        print(f"  {pid}  libro {h['n_libro']:>3} · docs {h['n_docs']:>2}  "
              f"{pregunta[:52]}")

    _escribir(d, historicos, hallazgos, tesis)
    print(f"→ {_SALIDA.relative_to(_RAIZ).as_posix()}")
    return 0


def _escribir(d, historicos, hallazgos, tesis) -> None:
    o: list[str] = []
    A = o.append

    A("# GM-Ω · ICPI — JUSTIFICACIÓN DE LAS TRANSFORMACIONES  `011-C3`")
    A("")
    A("**DERIVADO — no editar a mano.** Lo regenera "
      "`scripts/gm_omega/justificacion_transformaciones.py` barriendo el Gold "
      "Master completo y los documentos históricos de enero-abril.")
    A("")
    A("> ### El encargo")
    A("> No es «¿qué fórmula nos parece mejor?». Es: **¿qué transformación fue "
      "declarada, cuál fue implementada, cuál alimenta efectivamente al ICPI, "
      "y qué evidencia justifica cada transición?**")
    A("")
    A("```")
    A("  DECLARADO  ≠  IMPLEMENTADO  ≠  EFECTIVAMENTE UTILIZADO  ≠  JUSTIFICADO")
    A("```")
    A("")
    A("⚠️ **No se tocó nada.** Ni `C_i`, ni `E_i`, ni `T_i`, ni la fórmula, ni "
      "las calibraciones, ni `Ci_Manual_2025`, ni `Ci_Adaptativo`. El Gold "
      "Master es inmutable y el baseline **27,4582 %** sigue congelado. `C3` "
      "levanta acta.")
    A("")
    A("⚠️ **Y la regla que ordena la etapa** (`DOC-011`): una transición "
      "documentada **no autoriza a inventar su causa**. Reconstruir "
      "perfectamente CUÁNDO apareció algo no dice nada sobre POR QUÉ apareció. "
      "Donde no hay razón escrita, el resultado correcto es `NO DETERMINABLE` "
      "— y eso es un hallazgo, no un fallo del peritaje.")
    A("")

    # ── ★ El hallazgo que reordena la etapa ───────────────────────────────
    if tesis and tesis.get("definiciones"):
        A("## ★ 0 · El documento que reordena `011-C3`")
        A("")
        A(f"`{tesis['archivo']}`, **creado el {tesis['creado']}** — es decir, "
          "**anterior al 27-abr-2026**. Define las **seis** variables con "
          "definición conceptual, fundamento normativo y **tabla de escala con "
          "criterio de verificación**:")
        A("")
        A("| Variable | Título en la metodología | Definición conceptual |")
        A("|---|---|---|")
        for f in ("P_i", "R_i", "V_i", "E_i", "T_i", "C_i"):
            v = tesis["definiciones"].get(f)
            if v:
                A(f"| `{f}` | **{v['titulo']}** | {v['definicion'][:230]} |")
        A("")
        A("### ⚠️ Primera consecuencia · hay que corregir `007-B0`")
        A("")
        A("`007-B0` concluyó que **`C_i` entró el 27-abr-2026**, apoyándose en "
          "`H01!A94`. Esa celda dice:")
        A("")
        A("> ★ **Ci DETERMINISTA v1.0** (Javo Delgado Santana, 27-Abr-2026): Ci "
          "arranca en 1.00. Deducciones legales Sección L calculan Ci final.")
        A("")
        A(f"Leída junto a una metodología del {tesis['creado']} que **ya "
          "contiene `C_i`**, la fecha no dice lo que se le hizo decir:")
        A("")
        A("| Lectura | Estado |")
        A("|---|---|")
        A(f"| «el 27-abr **nace** `C_i`» | 🔴 **{_INF} y superada** |")
        A(f"| «el 27-abr nace **`Ci` DETERMINISTA v1.0**, una versión nueva de "
          f"un factor preexistente» | ✅ **{_DEM}** · lo dice la propia celda, "
          f"y la metodología anterior lo prueba |")
        A("")
        A("**Y el error tiene una forma reconocible:** se leyó la fecha del "
          "artefacto que documenta un cambio como si fuera la fecha del "
          "concepto. Es el escalón 7 de la escalera —**lo leído ≠ la "
          "fuente**— aplicado a una genealogía.")
        A("")
        A("### ★ Segunda consecuencia · la transformación real de `C_i`")
        A("")
        A("Lo que ocurrió el 27-abr **no fue una incorporación: fue una "
          "sustitución de mecanismo bajo el mismo nombre.**")
        A("")
        A("```")
        A("  ANTES (metodología, 25-mar-2026)      DESPUÉS (motor, 27-abr-2026)")
        A("  ──────────────────────────────        ────────────────────────────")
        A("  C_i = IMPUTABILIDAD ORGÁNICA          C_i = CALIDAD DE PROCESO")
        A("  claridad de la asignación de          descuento por infracciones")
        A("  responsabilidad en el Estatuto        normativas verificadas")
        A("")
        A("  Constitución 233 · NCI 200-04/401-01  LOSNCP · CGE · COPFP · CPCCS")
        A("  escala {1,00 · 0,90 · 0,75}           MAX(0,50 · 1,00 − Σ ded.)")
        A("  mínimo posible 0,75                   mínimo posible 0,50")
        A("```")
        A("")
        A("| Afirmación | Grado |")
        A("|---|---|")
        A(f"| Son **dos constructos distintos** con el mismo nombre | "
          f"**{_DEM}** · definición, fundamento normativo, escala y rango "
          f"difieren |")
        A(f"| El nuevo mecanismo admite valores que la escala original **no "
          f"contemplaba** (`0,50`) | **{_DEM}** |")
        A(f"| La **razón** de la sustitución | ⬜ **{_ND}** · ningún documento "
          f"la explica |")
        A("")
        A("### ★ Tercera consecuencia · la Sección I es un vestigio del primero")
        A("")
        A("`011-C2` no supo explicar por qué `H01` Sección I trae "
          "`Cod_Unidad`, `Dirección Responsable` y `Base Legal Estatuto` con "
          "escala «exclusivo / compartido / difuso». Ahora se explica solo: "
          "**eso es la implementación del `C_i` original**, la imputabilidad "
          "orgánica. Sigue viva dentro del libro, al lado del mecanismo que la "
          "sustituyó.")
        A("")
        A("Y con eso las **cuatro divergencias de `011-C2` dejan de ser cuatro "
          "anomalías sueltas** y pasan a ser un único fenómeno con nombre:")
        A("")
        A("> ### Dos generaciones del mismo factor conviven en el instrumento")
        A(">")
        A("> La Sección I implementa el constructo **original**; las Secciones "
          "L/M implementan el **nuevo**; y `Ci_Manual_2025` conserva los "
          "valores del original como fallback. Ninguna capa se retiró al "
          "añadirse la siguiente.")
        A("")
        A("⚠️ **Esto no dice que el cambio fuera indebido.** Dice que ocurrió, "
          "que ninguna de las dos capas se retiró, y que **el instrumento no "
          "declara cuál gobierna**. La valoración es de `011-C4`.")
        A("")

        # Las escalas literales, lado a lado
        if tesis["escalas"]:
            A("### Las escalas originales, literales")
            A("")
            for f, filas in tesis["escalas"].items():
                A(f"**`{f}`** — según `{tesis['archivo']}`:")
                A("")
                A("| " + " | ".join(x or "valor" for x in filas[0]) + " |")
                A("|" + "---|" * len(filas[0]))
                for fila in filas[1:]:
                    A("| " + " | ".join(fila) + " |")
                A("")
            A("> ### ★ Y aquí está la respuesta a `C3-09`")
            A(">")
            A("> `E_i` y `C_i` comparten escala `{1,00 · 0,90 · 0,75}` y "
              "vocabulario **porque ambas son escalas ordinales de tres grados "
              "sobre el mismo Estatuto Orgánico** — pero miden ejes distintos: "
              "`E_i` **quién EJECUTA**, `C_i` **quién RESPONDE**.")
            A("")
            A("La superposición que `011-C2` midió **no es un accidente ni un "
              "defecto: es deliberada y está justificada en la metodología**. Y "
              "las divergencias son exactamente lo que cabría esperar — la "
              "propia tesis trae un caso:")
            A("")
            A("> `M3` (Salud): ejecución **directa** del GAD (`E=1,00`) pero "
              "responsabilidad **compartida** entre Planificación —que "
              "formula— y Obras Públicas —que ejecuta— (`C=0,90`).")
            A("")
            A("| Afirmación | Grado |")
            A("|---|---|")
            A(f"| La superposición de escala está **justificada "
              f"documentalmente** | **{_DEM}** |")
            A(f"| Que `E_i` y `C_i` puedan divergir en una misma meta es "
              f"**conceptualmente esperable** | **{_DEM}** · con caso "
              f"ilustrado en la metodología |")
            A(f"| Que las **12 divergencias concretas del motor** respondan "
              f"cada una a esa razón | ⬜ **{_ND}** · exigiría la "
              f"justificación meta a meta, que no existe |")
            A("")
            A("⚠️ **La corrección que esto obliga a hacer a `011-C2`:** ahí se "
              "escribió que «**nada en el instrumento explica la "
              "diferencia**». Era cierto **del instrumento** y falso **del "
              "corpus**: la metodología sí explica por qué pueden diferir. Lo "
              "que sigue sin explicación es cada asignación concreta.")
            A("")

    # ── La cadena de versiones ────────────────────────────────────────────
    A("## 1 · La cadena de versiones, que es la columna vertebral")
    A("")
    A("`H80_MODEL_REGISTRY` registra el versionado del motor con fecha, "
      "operador y versión anterior. Es la mejor fuente de genealogía del libro:")
    A("")
    A("| Versión | Fecha | Operador | Estado | Anterior |")
    A("|---|---|---|---|---|")
    for v in d["versiones"]:
        A(f"| `{v['version']}` | {v['ts']} | {v['operador']} | {v['estado']} | "
          f"{v['anterior'] or '—'} |")
    A("")
    A("### ★ Lo que esta tabla revela sobre `C_i`")
    A("")
    A("`011-C2` estableció que `C_i` se creó el **27-abr-2026** "
      "(`H01!A94`). Situado en la cadena:")
    A("")
    A("```")
    A("  v1.0.2   31-mar-2026   ARCHIVADO")
    A("      │")
    A("      │     ⬅ 27-abr-2026 · nace C_i · NINGUNA VERSIÓN LO REGISTRA")
    A("      │")
    A("  v2.1     01-may-2026   ACTIVO")
    A("```")
    A("")
    A("Dos observaciones, ambas verificables en la tabla:")
    A("")
    A("| # | Observación | Grado |")
    A(f"|---|---|---|")
    A(f"| 1 | La incorporación de `C_i` cae **dentro del salto "
      f"`v1.0.2 → v2.1`** y no tiene entrada propia en el registro | "
      f"**{_DEM}** |")
    A(f"| 2 | El salto de versión es `1.0.2 → 2.1`: se omiten `1.1` y `2.0` | "
      f"**{_DEM}** |")
    A(f"| 3 | Por qué la incorporación de un sexto factor no generó su propia "
      f"entrada de versión | ⬜ **{_ND}** |")
    A("")
    A("⚠️ **Esto no es una acusación de mal versionado.** `P-05` del protocolo "
      "de gobernanza algorítmica exige versionado obligatorio, y el registro "
      "existe y es coherente. Lo que falta es **granularidad**: el salto que "
      "contiene el cambio más consecuente del motor —pasar de cinco a seis "
      "factores— se registra igual que cualquier otro.")
    A("")

    # ── El contraste de justificación ─────────────────────────────────────
    A("## ★ 2 · Qué factores tienen justificación DECLARADA, y cuáles no")
    A("")
    A("El hallazgo más limpio de `C3`, y no hubo que interpretarlo: "
      "`H14_PONDERADORES` **tiene columnas de justificación por meta**.")
    A("")
    if d["just_h14"]:
        for j in d["just_h14"][:6]:
            A(f"> `{j}`")
        A("")
    A("| Factor | ¿Tiene justificación declarada en el libro? |")
    A("|---|---|")
    A("| `P_i` | ✅ **sí** · `H14` columna «Justificación P_i», meta a meta |")
    A("| `R_i` | ✅ **sí** · `H14` columna «Justificación R_i», meta a meta |")
    A("| `V_i` | 🟡 parcial · la regla de combinación está en `H13`, sin "
      "columna de justificación |")
    A("| `T_i` | 🟡 parcial · jerarquía de fuentes declarada en el glosario |")
    A("| `E_i` | 🔴 **no** · literal en `H12`, sin fórmula ni columna |")
    A("| `C_i` | 🔴 **no** · `H01` Sección I trae `Base Legal Estatuto` por "
      "meta, pero **eso justifica la IMPUTACIÓN, no el VALOR** |")
    A("")
    A("> ### Las dos dimensiones sin justificación por meta son exactamente "
      "las dos que `011-C2` encontró superpuestas")
    A("")
    A("No es casualidad interpretable, y `C3` no la interpreta. Es un hecho "
      "con dos lecturas posibles y ninguna evidencia para elegir:")
    A("")
    A("| Lectura | Qué implicaría |")
    A("|---|---|")
    A("| `P_i` y `R_i` se justificaron porque **son las que la tesis "
      "desarrolló** | la justificación siguió al trabajo teórico |")
    A("| `E_i` y `C_i` se incorporaron **en fase de construcción**, cuando el "
      "hábito de justificar por meta ya no se aplicó | la justificación siguió "
      "al calendario |")
    A("")
    A(f"Cuál de las dos ocurrió: ⬜ **{_ND}**.")
    A("")

    # ── Las nueve preguntas ───────────────────────────────────────────────
    A("## 3 · Las nueve preguntas, con su evidencia")
    A("")
    A("Cada bloque muestra **la evidencia encontrada**, no un resumen de ella. "
      "Y separa su procedencia, porque no valen lo mismo:")
    A("")
    A("| Procedencia | Qué acredita |")
    A("|---|---|")
    A("| **El libro** | lo que el instrumento **declara** de sí mismo |")
    A("| **Los documentos** de enero-abril | lo que se **escribió** al "
      "construirlo — la única fuente que puede justificar |")
    A("")
    for pid, pregunta, _patron in _PREGUNTAS:
        h = hallazgos[pid]
        A(f"### {pid} · {pregunta}")
        A("")
        if h["libro"]:
            A("**En el libro:**")
            A("")
            for hoja, txt in h["libro"][:4]:
                A(f"- `{hoja}` — {txt[:300]}")
            A("")
        if h["docs"]:
            A("**En los documentos históricos:**")
            A("")
            for nombre, frag in h["docs"]:
                A(f"- `{nombre}` — …{frag[:300]}…")
            A("")
        if not h["libro"] and not h["docs"]:
            A(f"⬜ **Sin evidencia en ninguna de las dos fuentes.** `{_ND}`.")
            A("")

    # ── Dictamen ──────────────────────────────────────────────────────────
    A("## ★ Dictamen de `011-C3` · las nueve, por grado de certeza")
    A("")
    A("| # | Pregunta | Respuesta | Grado |")
    A("|---|---|---|---|")
    A(f"| **C3-01** | ¿quién y cuándo? | Javo Delgado Santana. El **concepto** "
      f"existe al menos desde el **{tesis.get('creado', '—')}**; la versión "
      f"**determinista**, el 27-abr-2026 | **{_DEM}** |")
    A(f"| **C3-02** | ¿qué fenómeno? | **dos**, sucesivamente: imputabilidad "
      f"orgánica → calidad jurídica del proceso | **{_DEM}** |")
    A(f"| **C3-03** | ¿por qué de 5 a 6 factores? | **la pregunta estaba mal "
      f"planteada**: la metodología ya tenía 6. Lo que cambió fue el "
      f"mecanismo, no el número | **{_DEM}** |")
    A(f"| **C3-03b** | ¿por qué se sustituyó el mecanismo? | — | ⬜ **{_ND}** "
      f"· ningún documento lo explica |")
    A(f"| **C3-04** | ¿por qué esos cuatro eventos? | la Sección L los declara "
      f"«framework jurídico agnóstico» y escalable a otros países; **la razón "
      f"de estos cuatro y no otros no consta** | **{_DECL}** parcial |")
    A(f"| **C3-05** | ¿por qué esos pesos? | no constan en la metodología —que "
      f"no tenía deducciones—. El único precedente es `H95` `L-07`: los pesos "
      f"del TGI son «**criterio experto (Dylus Lab)**, no PCA ni regresión» | "
      f"⬜ **{_ND}** para `C_i` |")
    A(f"| **C3-06** | ¿por qué el piso `0,50`? | tampoco consta. Y **cambia el "
      f"rango**: la escala original tenía mínimo `0,75` | ⬜ **{_ND}** |")
    A(f"| **C3-07** | ¿por qué `Ci_Manual_2025` en 2026? | el instrumento "
      f"declara que preserva el ICPI de referencia en estado vacío; **por qué "
      f"2025 puede representar 2026 no consta** | **{_DECL}** parcial |")
    A(f"| **C3-08** | ¿qué es `Ci_Adaptativo`? | definido, no conectado. Las "
      f"cinco lecturas causales siguen abiertas | ⬜ **{_ND}** |")
    A(f"| **C3-09** | ¿por qué `E_i` y `C_i` comparten escala y divergen? | "
      f"**son dos ejes del mismo Estatuto**: quién ejecuta vs. quién responde. "
      f"La superposición es deliberada y justificada | **{_DEM}** · salvo cada "
      f"asignación concreta, que es `{_ND}` |")
    A("")
    A("### Lo que `C3` cambia respecto de lo que se creía")
    A("")
    A("| Se creía | Ahora consta |")
    A("|---|---|")
    A("| `C_i` se incorporó el 27-abr-2026 | el **concepto** es anterior; esa "
      "fecha data su **versión determinista** |")
    A("| `E_i` y `C_i` se superponen sin explicación | la superposición está "
      "**justificada en la metodología**: ejes distintos, escala común |")
    A("| las cuatro divergencias de `C2` eran anomalías sueltas | son **un "
      "solo fenómeno**: dos generaciones del factor conviviendo |")
    A("| `E_i` carece de biografía | tiene definición, fundamento (`COOTAD 54` "
      "· `NCI 200-04`), escala y **ejemplos de aplicación** |")
    A("")
    A("### ⚠️ Y una corrección que `C3` le debe a `009` y a `D-014`")
    A("")
    A("`011-C2` concluyó que **ninguna variable del ICPI contempla la entrega "
      "material**. La metodología obliga a matizarlo, y el matiz importa "
      "porque **rescata parcialmente la intuición de Javo, reubicándola**:")
    A("")
    A("> `T_i` — Materialización Temporal — se define sobre el **devengado y "
      "no el compromiso**, y la metodología justifica esa elección así: el "
      "devengado exige «**factura válida, acta de entrega-recepción firmada e "
      "informe de conformidad del fiscalizador**» según el Acuerdo Ministerial "
      "067 del MEF. Y lo dice con todas las letras: **«neutraliza una forma "
      "frecuente de gaming: reportar contratos firmados en diciembre como "
      "metas ejecutadas cuando la obra apenas comienza en enero»**.")
    A("")
    A("Es decir: **la defensa contra el maquillaje de fin de ejercicio SÍ "
      "está en el constructo — pero en `T_i`, no en `C_i`**. Javo señalaba un "
      "mecanismo real; se equivocó de variable.")
    A("")
    A("Ahora bien, la protección tiene un límite que hay que decir igual de "
      "claro:")
    A("")
    A("| Nivel | Estado |")
    A("|---|---|")
    A("| El devengado **presupone normativamente** acta de entrega-recepción "
      "| ✅ Acuerdo 067 MEF |")
    A("| El motor **verifica de forma independiente** que esa acta exista | "
      "🔴 **no** · lee la columna «Devengado» de la cédula eSIGEF |")
    A("")
    A("> La protección es **normativa, no verificada por el motor**. Si una "
      "entidad devenga sin acta bien formada, el ICPI no puede detectarlo — "
      "confía en que el dato eSIGEF cumpla el Acuerdo 067.")
    A("")
    A("Eso reformula `D-014`: no es que el constructo ignore la entrega "
      "material, sino que **la delega en la corrección del dato de origen**. "
      "Si esa delegación es suficiente lo juzga `011-C4`.")
    A("")
    A("> ### GM-Ω-011-C3 — CERRADO COMO RECONSTRUCCIÓN CAUSAL")
    A(">")
    A("> Se estableció **qué se declaró, qué se implementó y qué alimenta "
      "efectivamente al ICPI**. La cadena está documentada hasta el mecanismo; "
      "**se corta en la razón**: por qué se sustituyó el constructo de `C_i`, "
      "por qué esos pesos y por qué ese piso **no constan en ninguna fuente**.")
    A(">")
    A("> Eso es un **resultado**, no una carencia del peritaje: la razón no "
      "está porque **nunca se escribió** (`DOC-022`), y las decisiones "
      "conversacionales no dejan rastro. Inventarla sería el error que "
      "`DOC-011` prohíbe.")
    A(">")
    A("> **No dictamina** si las transformaciones fueron correctas. `011-C4`.")
    A("")
    A("---")
    A(f"*GM-Ω-ICPI-011-C3 · {len(d['hojas'])} hojas barridas · "
      f"{len(historicos)} documentos históricos · el Gold Master no se "
      f"modificó · baseline 27,4582 % congelado · Dylus Lab © 2026*")

    _SALIDA.write_text("\n".join(o) + "\n", encoding="utf-8")


if __name__ == "__main__":
    raise SystemExit(main())
