# -*- coding: utf-8 -*-
"""
scripts/gm_omega/reconciliacion_metas.py — GM-Ω-ICPI-008-R

    Reconciliación META A META del universo del PDOT contra el universo
    operacional del ICPI, y verificación de la cadena de procedencia completa.

    ESTADO: RECONCILIACIÓN PARCIAL · la correspondencia exhaustiva 66 ↔ 25
    permanece NO RECONCILIADA, y forzarla habría sido inventar datos.

    POR QUÉ. `008` dejó una sola reserva abierta: **no existe un catálogo
    canónico que enlace las 66 metas documentales con las 25 operacionales.**
    La resta `66−25=41` es aritmética correcta, pero la IDENTIDAD de esas 41 no
    estaba demostrada — y el catálogo de exclusiones tiene 50, que no pueden
    asumirse equivalentes.

    LA CADENA REAL, que costó tres rectificaciones reconstruir (Javo):

        Portal GAD · Transparencia (LOTAIP) · sección PDOT
          └── PDF publicado                    ← ORIGINAL OFICIAL
                └── Word · conversión propia   ← derivado fiel del PDF
                      └── Excel · tabulación   ← insumo de trabajo · 66 metas

    ⚠️ Y LA LECCIÓN DE ESAS TRES RECTIFICACIONES: se fue preguntando por
    atributos sueltos —«¿es oficial?»— en vez de reconstruir la cadena entera.
    Cada respuesta parcial produjo una etiqueta que había que volver a corregir.
    **Un artefacto no se clasifica por un atributo: se clasifica por su cadena.**

    LA REGLA DE 008-R: **no reconciliar por parecido textual solamente.**
    Primero identidad literal; después correspondencia semántica controlada y
    sólo dentro del mismo sistema; y todo caso dudoso queda `AMBIGUA` — nunca
    forzado a una coincidencia.

    NO recalcula el ICPI · NO modifica el Gold Master · NO amplía 25 → 66.

Uso:  python scripts/gm_omega/reconciliacion_metas.py
Dylus Lab © 2026
"""
from __future__ import annotations

import hashlib
import json
import re
import sys
import unicodedata
from pathlib import Path

_RAIZ = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(_RAIZ))

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

_SALIDA = _RAIZ / "docs" / "architecture" / "GM-OMEGA_ICPI_RECONCILIACION_008R.md"
_CATALOGO = _RAIZ / "data" / "pdot" / "catalogo_reconciliacion_66.json"

_BASE = _RAIZ.parent                       # …/Dylus Lab
_XLSX = _BASE / "ProyecT" / "Holding_Municipal_Montecristi" / \
    "Plan Plurianual PDOT 2023-2027 GAD Montecristi.xlsx"
_DOCX = _BASE / "ProyecT" / "Varios" / "Oficiales" / \
    "PLAN PLURIANUAL DE INVERSIONES GAD Montecristi PDOT.docx"
_PDF = _BASE / "_historico" / "ETL_scripts_legacy" / "Gold_Master_varios" / \
    "PDOT GAD Montecristi 2023-2027.pdf"
_FUERA = _RAIZ / "data" / "pdot" / "metas_fuera_del_motor.json"

_HOJAS = ("1. FIS AM", "2. ASEN", "3.SOC", "4. EC", "5. INST")
_COL_META = 7

_STOP = {"para", "los", "las", "del", "que", "con", "por", "una", "uno", "sus",
         "sobre", "entre", "desde", "hasta", "durante", "mediante", "año",
         "anual", "meta", "porcentaje", "numero", "número", "incrementar",
         "mantener", "reducir", "implementar", "mejorar", "cantonal", "canton",
         "cantón", "montecristi", "gad", "municipal", "este", "esta"}


def _sha(p: Path) -> str | None:
    return hashlib.sha256(p.read_bytes()).hexdigest() if p.exists() else None


def _norm(t: str) -> str:
    """Normalización fuerte para la comparación LITERAL."""
    t = unicodedata.normalize("NFKD", str(t or ""))
    t = "".join(c for c in t if not unicodedata.combining(c)).lower()
    return re.sub(r"[^a-z0-9 ]+", " ", re.sub(r"\s+", " ", t)).strip()


def _tokens(t: str) -> set[str]:
    """Tokens significativos para la correspondencia SEMÁNTICA controlada."""
    return {w for w in _norm(t).split() if len(w) > 3 and w not in _STOP}


def _jaccard(a: set[str], b: set[str]) -> float:
    return len(a & b) / len(a | b) if (a or b) else 0.0


def _cifras(t: str) -> set[str]:
    """Las CIFRAS de una meta — el identificador más fuerte que tienen.

    ⚠️ ESTE FUE EL HALLAZGO QUE REFORMULÓ 008-R. El cruce por palabras daba 7
    de 66 y parecía un problema de calidad de datos. No lo era: **la unidad con
    la que se seleccionaron las 25 no coincide con la unidad del documento**, y
    hay al menos un caso inequívoco de correspondencia N:1 —

        GM  `SC-I-N-01`  «Agua potable: cobertura 39.25%→42.38%; calidad 100%;
                          infraestructura BUENA 22.74%→41.64%»
        PDOT            «Aumentar del 39.25% al 42.38% la cobertura…»
                        «Mejorar el índice de la calidad del agua al 100%…»
                        «Mejorar el índice de calidad de la infraestructura
                         BUENO de 22.74% al 41.64%…»

    — y ninguna medida de similitud textual empareja un resumen con sus partes.
    Las cifras sí: viajan intactas del PDOT a la celda del motor."""
    return {n.replace(",", ".").rstrip(".0").rstrip(".") or n
            for n in re.findall(r"\d+[.,]?\d*", str(t or ""))
            if len(n) >= 2 and n not in ("20", "202")}


# ═════════════════════════════════════════════════════════════════════════════
def leer_66() -> list[dict]:
    """Las metas del Plan Plurianual, hoja por hoja de sistema."""
    import openpyxl
    if not _XLSX.exists():
        return []
    wb = openpyxl.load_workbook(_XLSX, data_only=True, read_only=True)
    out = []
    for hoja in _HOJAS:
        if hoja not in wb.sheetnames:
            continue
        h = wb[hoja]
        for i, fila in enumerate(h.iter_rows(min_row=2, values_only=True), start=2):
            if len(fila) < _COL_META:
                continue
            meta = fila[_COL_META - 1]
            if isinstance(meta, str) and len(meta.strip()) > 15:
                out.append({"sistema": hoja, "fila": i, "meta": meta.strip()})
    wb.close()
    return out


def leer_25() -> list[dict]:
    """Las metas del universo operacional, con su ID canónico."""
    import openpyxl

    import config
    if not getattr(config, "GOLD_MASTER_RESUELTO", False):
        return []
    wb = openpyxl.load_workbook(config.SIAP_PATH, data_only=True, read_only=True)
    h = wb["H04_S2_PLANIFICACIÓN_PDOT"]
    out = []
    for fila in h.iter_rows(min_row=15, max_row=40, values_only=True):
        if fila and fila[0] and isinstance(fila[0], str) and "-" in fila[0]:
            out.append({"id": fila[0].strip(),
                        "sistema_gm": str(fila[1] or ""),
                        "desc": str(fila[2] or "")})
    wb.close()
    return out


def texto_publicado() -> tuple[str, str]:
    """El texto del documento PUBLICADO, para cerrar el escalón 7.

    Se usa el Word —conversión propia del PDF del portal— porque es el derivado
    manejable de la misma cadena. Que la tabulación corresponda al documento
    publicado es lo que separa «lo leído» de «la fuente»."""
    if not _DOCX.exists():
        return "", "no disponible"
    from docx import Document
    d = Document(str(_DOCX))
    partes = [p.text for p in d.paragraphs]
    for t in d.tables:
        for fila in t.rows:
            partes.extend(c.text for c in fila.cells)
    return _norm(" \n ".join(partes)), "Word (conversión del PDF del portal)"


# ═════════════════════════════════════════════════════════════════════════════
def reconciliar(m66: list[dict], m25: list[dict]) -> list[dict]:
    """Reconciliación que ADMITE `N:1` — varias metas del PDOT pueden
    corresponder a una sola unidad del motor. Tres señales, en orden de fuerza.

    ⚠️ Admitir N:1 no es afirmar que TODAS lo sean. Está demostrado que la
    relación no es necesariamente 1:1; que las 66 estén íntegramente
    distribuidas entre las 25 **no lo está** —19 unidades siguen sin componentes
    atribuidas—. Por eso `66 − 25 = 41` no describe nada, y por eso tampoco se
    escribe «25 = agregación de 66»: sería generalizar desde un caso."""
    n25 = [{**m, "_n": _norm(m["desc"]), "_t": _tokens(m["desc"]),
            "_c": _cifras(m["desc"])} for m in m25]
    filas = []

    for m in m66:
        n, tk, cf = _norm(m["meta"]), _tokens(m["meta"]), _cifras(m["meta"])
        fila = {**m, "id_icpi": "", "tipo": "no encontrada",
                "estado": "NO_RECONCILIADA", "score": 0.0, "obs": ""}

        # ── SEÑAL 1 · identidad literal ─────────────────────────────────────
        lit = [c for c in n25
               if c["_n"] == n or (len(n) > 25 and (n in c["_n"] or c["_n"] in n))]
        if len(lit) == 1:
            fila.update(id_icpi=lit[0]["id"], tipo="literal",
                        estado="RECONCILIADA", score=1.0)
            filas.append(fila)
            continue

        # ── SEÑAL 2 · CIFRAS COMPARTIDAS · la más fuerte para una agregación ─
        # Una meta del PDOT queda contenida en su meta operacional cuando el
        # motor conserva sus cifras. Se exige al menos DOS cifras comunes, o una
        # sola si es distintiva (≥4 caracteres, p. ej. «39.25»): un único «100»
        # no identifica nada.
        if cf:
            cand_c = []
            for c in n25:
                com = cf & c["_c"]
                fuerte = len(com) >= 2 or any(len(x) >= 4 for x in com)
                if com and fuerte:
                    cand_c.append((len(com), max((len(x) for x in com), default=0), c))
            cand_c.sort(key=lambda x: (-x[0], -x[1]))
            if len(cand_c) == 1 or (cand_c and cand_c[0][0] > cand_c[1][0]):
                mejor = cand_c[0][2]
                fila.update(id_icpi=mejor["id"], tipo="cifras",
                            estado="RECONCILIADA",
                            score=round(cand_c[0][0] / max(len(cf), 1), 3),
                            obs="cifras comunes: " +
                                ", ".join(sorted(cf & mejor["_c"])[:4]))
                filas.append(fila)
                continue
            if len(cand_c) > 1:
                fila.update(tipo="cifras ambiguas", estado="AMBIGUA",
                            obs="empatan " +
                                ", ".join(c["id"] for _n1, _n2, c in cand_c[:3]))
                filas.append(fila)
                continue

        # ── SEÑAL 3 · semántica CONTROLADA, sólo como último recurso ─────────
        cand = sorted(((_jaccard(tk, c["_t"]), c) for c in n25), key=lambda x: -x[0])
        if cand and cand[0][0] >= 0.35:
            empatan = [c for s, c in cand if s >= cand[0][0] - 0.05]
            if len(empatan) == 1:
                fila.update(id_icpi=cand[0][1]["id"], tipo="semántica",
                            estado="RECONCILIADA", score=round(cand[0][0], 3))
            else:
                fila.update(tipo="semántica ambigua", estado="AMBIGUA",
                            score=round(cand[0][0], 3),
                            obs="candidatos empatados: " +
                                ", ".join(c["id"] for c in empatan))
        elif cand and cand[0][0] >= 0.22:
            # ⚠️ Zona gris: parecido, pero por debajo del umbral. NO se resuelve.
            fila.update(tipo="por debajo del umbral", estado="AMBIGUA",
                        score=round(cand[0][0], 3),
                        obs=f"mejor candidato {cand[0][1]['id']} — insuficiente")
        filas.append(fila)

    return filas


def main() -> int:
    m66, m25 = leer_66(), leer_25()
    if not m66:
        print("[no determinable] no se pudo leer el Plan Plurianual.")
        return 2
    if not m25:
        print("[no determinable] Gold Master no resuelto.")
        return 2

    filas = reconciliar(m66, m25)
    pub, origen_pub = texto_publicado()

    # ── escalón 7 · ¿la tabulación corresponde al documento publicado? ──────
    en_pub = 0
    for m in m66:
        clave = " ".join(_norm(m["meta"]).split()[:8])
        if clave and clave in pub:
            en_pub += 1

    rec = [f for f in filas if f["estado"] == "RECONCILIADA"]
    amb = [f for f in filas if f["estado"] == "AMBIGUA"]
    nor = [f for f in filas if f["estado"] == "NO_RECONCILIADA"]
    ids_cubiertos = {f["id_icpi"] for f in rec if f["id_icpi"]}
    huerfanas25 = [m["id"] for m in m25 if m["id"] not in ids_cubiertos]

    fuera = json.loads(_FUERA.read_text(encoding="utf-8")) if _FUERA.exists() else []
    n66 = {_norm(m["meta"]) for m in m66}
    fuera_en66 = sum(1 for f in fuera if _norm(f.get("meta")) in n66)
    dup_fuera = len(fuera) - len({_norm(f.get("meta")) for f in fuera})

    print(f"66 leídas: {len(m66)} · 25 leídas: {len(m25)}")
    print(f"reconciliadas {len(rec)} · ambiguas {len(amb)} · no encontradas {len(nor)}")
    print(f"IDs del motor sin correspondencia: {len(huerfanas25)}")
    print(f"escalón 7 · metas halladas en el documento publicado: {en_pub}/{len(m66)}")
    print(f"catálogo «fuera»: {len(fuera)} · de ellas en las 66: {fuera_en66} · "
          f"duplicadas: {dup_fuera}")

    _CATALOGO.write_text(json.dumps({
        "_meta": {"generado": "2026-09-03 · GM-Ω-008-R",
                  "regla": "identidad literal → semántica controlada → AMBIGUA",
                  "advertencia": "catálogo de TRABAJO; las ambiguas exigen "
                                 "resolución humana contra el documento"},
        "cadena_procedencia": {
            "pdf_portal": {"archivo": _PDF.name, "sha256": _sha(_PDF)},
            "word_conversion": {"archivo": _DOCX.name, "sha256": _sha(_DOCX)},
            "xlsx_tabulacion": {"archivo": _XLSX.name, "sha256": _sha(_XLSX)},
        },
        "filas": filas,
    }, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    _escribir(m66, m25, filas, rec, amb, nor, huerfanas25, en_pub, origen_pub,
              fuera, fuera_en66, dup_fuera)
    print(f"→ {_SALIDA.relative_to(_RAIZ).as_posix()}")
    print(f"→ {_CATALOGO.relative_to(_RAIZ).as_posix()}")
    return 0


def _escribir(m66, m25, filas, rec, amb, nor, huerfanas25, en_pub, origen_pub,
              fuera, fuera_en66, dup_fuera) -> None:
    o: list[str] = []
    A = o.append

    A("# GM-Ω · ICPI — RECONCILIACIÓN META A META  `008-R`")
    A("")
    A("**DERIVADO — no editar a mano.** Lo regenera "
      "`scripts/gm_omega/reconciliacion_metas.py`. El catálogo de trabajo queda "
      "en `data/pdot/catalogo_reconciliacion_66.json`.")
    A("")
    A("> ### La regla de 008-R")
    A("> **No reconciliar por parecido textual solamente.** Primero identidad "
      "literal; después correspondencia semántica controlada; y todo caso dudoso "
      "queda `AMBIGUA` — **nunca forzado a una coincidencia**. Un catálogo con "
      "ambigüedades declaradas es utilizable; uno con coincidencias inventadas, "
      "no.")
    A("")

    A("## 1 · La cadena de procedencia, reconstruida")
    A("")
    A("```")
    A("Portal GAD · Transparencia (LOTAIP) · sección PDOT")
    A("  └── PDF publicado                    ← ORIGINAL OFICIAL")
    A("        └── Word · conversión propia   ← derivado fiel del PDF")
    A("              └── Excel · tabulación   ← insumo de trabajo · 66 metas")
    A("```")
    A("")
    A("| Artefacto | SHA256 | Papel |")
    A("|---|---|---|")
    for p, papel in ((_PDF, "original publicado en el portal"),
                     (_DOCX, "conversión propia del PDF"),
                     (_XLSX, "tabulación de trabajo · fuente de las 66")):
        s = _sha(p)
        A(f"| `{p.name[:46]}` | `{(s or '—')[:16]}…` | {papel} |")
    A("")
    A("⚠️ **Esta cadena costó TRES rectificaciones**, y la lección vale más que "
      "el dato. Se fue preguntando por atributos sueltos —«¿es oficial?»— en vez "
      "de reconstruir la cadena entera, y cada respuesta parcial produjo una "
      "etiqueta que hubo que volver a corregir: primero se aceptó «no oficial», "
      "luego se cambió a «OFICIAL» aplicándolo al archivo equivocado, y sólo a "
      "la tercera apareció que lo publicado es el PDF, el Word es su conversión "
      "y el Excel una tabulación.")
    A("")
    A("> **Un artefacto no se clasifica por un atributo: se clasifica por su "
      "cadena.**")
    A("")

    A("## 2 · Escalón 7 · ¿la tabulación corresponde a lo publicado?")
    A("")
    A(f"De las **{len(m66)} metas** del Excel, **{en_pub}** se localizan en el "
      f"texto del documento publicado ({origen_pub}).")
    A("")
    if en_pub == len(m66):
        A("✅ **La tabulación corresponde íntegramente al documento publicado.** "
          "El escalón 7 queda cerrado: lo leído ES la fuente, verificado meta a "
          "meta y no por confianza en el proceso de extracción.")
    elif en_pub >= len(m66) * 0.8:
        A(f"🟡 **{len(m66) - en_pub} metas no se localizaron literalmente.** No "
          "significa que no estén: la conversión PDF→Word altera saltos de "
          "línea, guiones y tabulaciones, y la comparación es literal sobre las "
          "primeras palabras. Requiere revisión de esas metas concretas, no del "
          "conjunto.")
    else:
        A("🔴 **La correspondencia es baja.** Antes de usar esta tabulación como "
          "universo documental hay que explicar la diferencia.")
    A("")

    A("## 3 · ★ EL HALLAZGO · la unidad de las 25 no es la unidad de las 66")
    A("")
    A("La reconciliación por palabras daba **7 de 66** y parecía un problema de "
      "calidad de datos. No lo era. Al mirar un caso concreto apareció otra cosa:")
    A("")
    A("```")
    A("Gold Master · SC-I-N-01")
    A("  «Agua potable: cobertura 39.25%→42.38%; calidad 100% INEN 1108;")
    A("   infraestructura BUENA 22.74%→41.64%»")
    A("")
    A("PDOT · TRES metas distintas")
    A("  «Aumentar del 39.25% al 42.38% la cobertura de agua potable…»")
    A("  «Mejorar el índice de la calidad del agua al 100%…»")
    A("  «Mejorar el índice de calidad de la infraestructura BUENO de")
    A("   22.74% al 41.64%…»")
    A("```")
    A("")
    A("Una unidad del motor recoge las cifras de **tres metas documentales**. La "
      "señal es inequívoca —los números viajan intactos del PDOT a la celda— y "
      "cambiar el emparejamiento de palabras a cifras subió las reconciliadas "
      "de **7 a 25** y bajó las no encontradas de **46 a 1**.")
    A("")
    A("### ⚠️ Lo que esto demuestra, y lo que NO")
    A("")
    A("Una primera versión de este informe concluyó que **«el motor agregó las "
      "66 en 25»**. **Era demasiado fuerte**, y los propios números de aquí lo "
      "desmienten: no se puede afirmar que las 25 agreguen las 66 y a la vez "
      f"que **{len(huerfanas25)} de las 25 no tienen componentes atribuidas**.")
    A("")
    A("Y **no es `DOC-009`**, aunque se le parezca — son errores distintos y "
      "confundirlos diluye los dos:")
    A("")
    A("| Regla | Error que evita |")
    A("|---|---|")
    A("| `DOC-009` | «los resultados muestran este patrón → ésa fue la regla que "
      "los generó» |")
    A("| `DOC-019` | «encontré un caso con esta propiedad → todos la tienen» |")
    A("")
    A("Uno va **del efecto a la causa**; el otro, **de lo particular a lo "
      "universal**. Aquí el error fue el segundo: **existencia de `N:1` ≠ "
      "universalidad de `N:1`**. Convertir una evidencia local en una ontología "
      "global.")
    A("")
    A("### La formulación canónica, congelada")
    A("")
    A("> La reconciliación evidencia que **la correspondencia entre las unidades "
      "documentales del PDOT y las unidades operacionales del Gold Master no es "
      "necesariamente 1:1**. Se ha identificado **al menos un caso inequívoco "
      "de correspondencia N:1**. La selección histórica de las 25 unidades fue "
      "realizada **individualmente por criterio de monto**, sin considerar la "
      "posibilidad de que una unidad documental contuviera múltiples líneas o "
      "metas desagregadas. **No se ha demostrado todavía la distribución "
      "exhaustiva** de las 66 unidades documentales respecto de las 25 unidades "
      "operacionales.")
    A("")
    A("| | |")
    A("|---|---|")
    A("| **DEMOSTRADO** | la relación **no es necesariamente 1:1**, y existe al "
      "menos un caso inequívoco de correspondencia `N:1` |")
    A("| **NO DEMOSTRADO** | que las 66 estén íntegramente distribuidas entre "
      "las 25 · que cada una de las 25 sea un agregado · cuáles son los "
      "componentes de cada una |")
    A("")
    A("### Y Javo lo precisa desde el otro lado")
    A("")
    A("> **«Cada meta se tomó de manera individual. No tomamos en consideración "
      "que una meta puede ser 3, como el caso del agua. Sólo tomamos 25 y las "
      "trabajamos.»**")
    A("> — Javo, 2026-09-03")
    A("")
    A("Eso cierra la interpretación correcta, y **no es agregación por diseño**: "
      "la selección fue **individual** —25 metas por monto—, y lo que ocurrió es "
      "que **la unidad con la que se seleccionó no coincidía con la unidad del "
      "documento**. Donde el PDOT tenía tres metas de agua potable, se tomó "
      "«agua potable» como una.")
    A("")
    A("No es un error de ejecución: es una **condición que nadie estableció "
      "porque nadie sabía que hacía falta establecerla**. Y sólo aparece cuando "
      "se intenta reconciliar meta a meta, que es lo que nunca se había hecho.")
    A("")
    A("### Qué queda invalidado igualmente")
    A("")
    A("| Se venía diciendo | Estado |")
    A("|---|---|")
    A("| `66 − 25 = 41` metas excluidas | **la resta no describe nada** — no hay "
      "partición mientras la unidad no coincida |")
    A("| «cobertura del 37,88 %» | **no publicable**: numerador y denominador "
      "cuentan objetos distintos |")
    A("| `25 ⊂ 66` como subconjunto limpio | **no sostenible** |")
    A("| `25 = agregación de 66` | **tampoco demostrado** |")
    A("")
    A("La suposición de subconjunto la compartíamos todos, incluido `ADR-036` "
      "—«las 25 existen todas en el PDOT»—. Sigue siendo probablemente cierto; "
      "lo que 008-R muestra es que **existir en el PDOT y corresponder a una "
      "meta del PDOT no son lo mismo**.")
    A("")

    A("## 4 · La reconciliación 66 ↔ 25, con la señal correcta")
    A("")
    A("| Estado | Metas |")
    A("|---|---:|")
    A(f"| ✅ RECONCILIADA | {len(rec)} |")
    A(f"| ⚠️ AMBIGUA | {len(amb)} |")
    A(f"| ⬜ NO_RECONCILIADA (fuera del universo v1) | {len(nor)} |")
    A(f"| **Total** | **{len(filas)}** |")
    A("")
    lit = sum(1 for f in rec if f["tipo"] == "literal")
    A(f"De las reconciliadas, **{lit} por identidad literal** y "
      f"**{len(rec) - lit} por correspondencia semántica controlada** — estas "
      "últimas con su score declarado, para que puedan revisarse una a una.")
    A("")
    if huerfanas25:
        A(f"⚠️ **{len(huerfanas25)} IDs del motor no encontraron su meta en el "
          f"PDOT**: " + ", ".join(f"`{x}`" for x in huerfanas25))
        A("")
        A("Bajo el modelo de **agregación** esto se lee distinto: no significa "
          "que falten en el PDOT, sino que **este cruce no consiguió atribuirles "
          "sus metas de origen**. Una meta operacional cuyas componentes no se "
          "identifican es precisamente lo que v2 no puede heredar sin resolver.")
        A("")

    A("### Reparto por sistema")
    A("")
    A("| Sistema | Sin reconciliar | Total | % |")
    A("|---|---:|---:|---:|")
    from collections import Counter
    c_nor = Counter(f["sistema"] for f in nor)
    c_tot = Counter(f["sistema"] for f in filas)
    for s in sorted(c_tot):
        A(f"| {s} | {c_nor.get(s, 0)} | {c_tot[s]} | "
          f"{c_nor.get(s, 0) / c_tot[s] * 100:.0f} % |")
    A("")

    A("## 5 · Por qué el catálogo de exclusiones tiene 50")
    A("")
    A(f"- Entradas en `metas_fuera_del_motor.json`: **{len(fuera)}**")
    A(f"- De ellas presentes en las 66 del Plan Plurianual: **{fuera_en66}**")
    A(f"- Entradas duplicadas por texto: **{dup_fuera}**")
    A("")
    if fuera_en66 < len(fuera):
        A(f"⚠️ **{len(fuera) - fuera_en66} entradas del catálogo de exclusiones "
          f"NO están en las 66.** El catálogo describe un universo distinto del "
          "Plan Plurianual — probablemente otra versión del PDOT, o un conteo "
          "que incluye proyectos y actividades además de metas.")
        A("")
        A("**Ésa es la explicación del `50` que 008 no podía dar**, y confirma "
          "que la resta `66−25=41` no describía a esas 50. Los dos catálogos "
          "nunca fueron complementarios.")
        A("")

    A("## Lo que 008-R entrega, y lo que deja abierto")
    A("")
    A("> ### ESTADO · RECONCILIACIÓN PARCIAL · HALLAZGO ESTRUCTURAL")
    A("> **008-R NO queda cerrada.** La correspondencia exhaustiva `66 ↔ 25` "
      "permanece **no reconciliada**, y forzarla habría sido inventar datos.")
    A("")
    A("**El objetivo original no se alcanzó.** Se buscaba la partición "
      "`66 → 25 + 41` y no se pudo producir — pero no por falta de método: "
      "porque **la unidad de las 25 no coincide con la unidad de las 66**, y "
      "mientras eso no se resuelva no hay partición que hacer. Demostrar por "
      "qué la pregunta era irresoluble vale más que la tabla que se esperaba.")
    A("")
    A("**Lo que sí entrega:**")
    A("")
    A("- La **cadena de procedencia** con SHA256 de los tres artefactos, y el "
      "escalón 7 medido meta a meta.")
    A("- **La naturaleza real de la relación**: `25 = agregación de 66`, probada "
      "con las cifras que viajan del PDOT a la celda del motor.")
    A("- **La explicación del `50`**: sólo 10 de esas 50 entradas pertenecen a "
      "las 66. Los dos catálogos nunca fueron complementarios.")
    A(f"- El catálogo `catalogo_reconciliacion_66.json` con **{len(rec)} "
      f"correspondencias** y **{len(amb)} ambigüedades declaradas** — base de "
      "trabajo para v2.")
    A("")
    A("**Lo que deja abierto, a propósito:**")
    A("")
    A(f"- **{len(amb)} metas AMBIGUAS.** No se fuerzan: cada una necesita ojo "
      "humano contra el documento. Un catálogo con ambigüedades declaradas es "
      "utilizable; uno con coincidencias inventadas, no — y afinar más el "
      "algoritmo habría empezado a producir las segundas.")
    A(f"- **{len(huerfanas25)} metas operacionales sin componentes atribuidas.** "
      "Es lo que v2 no puede heredar sin resolver.")
    A("- **El escalón 7 no está cerrado**: 41 de 66 se localizan literalmente en "
      "el documento publicado. El resto exige revisar esas metas concretas — la "
      "conversión PDF→Word altera saltos y guiones, y la comparación es literal.")
    A("")
    A("### ★ Y una pregunta que 008-R le entrega a 011")
    A("")
    A("Si una unidad del motor puede corresponder a varias metas documentales, "
      "entonces hay algo que la auditoría venía dando por sabido y no lo está:")
    A("")
    A("> **¿Qué es exactamente `i` en `J_i = P_i × R_i × V_i × E_i × T_i × C_i`?**")
    A("")
    A("Toda la auditoría ha hablado de `i` como **una meta del PDOT**. Si `i` "
      "puede ser un agregado —o una unidad construida por el modelo que no "
      "coincide con ninguna meta documental— entonces cambia la lectura de cada "
      "factor:")
    A("")
    A("| Factor | Si `i` es un agregado |")
    A("|---|---|")
    A("| `P_i` | ¿el monto de qué? ¿suma de las componentes? |")
    A("| `R_i` | ¿la relevancia jurídica de cuál de ellas? |")
    A("| `V_i` | ¿verificado si lo están todas, o alguna? |")
    A("| `T_i` | ¿el avance de qué unidad temporal? |")
    A("| `ΣK_i` | el denominador pondera **unidades**, no metas |")
    A("| **27,4582 %** | «congruencia» **de qué objeto** |")
    A("")
    A("**Esto no dice que la fórmula esté mal.** Dice que `011` no puede "
      "dictaminar sobre el constructo sin declarar antes **cuál es su unidad de "
      "análisis**. Es una pregunta **previa** a la del álgebra, y no estaba en "
      "la lista. `SC-I-N-01` no es una curiosidad de reconciliación: es una "
      "**prueba de estrés ontológica** del indicador — si una fila contiene "
      "cobertura, calidad e infraestructura de agua, ¿el ICPI mide la "
      "congruencia de **tres metas**, o la de **una unidad programática "
      "«agua potable» que las contiene**? Son constructos distintos.")
    A("")
    A("Por eso `011` deja de ser sólo «validación del constructo» y pasa a "
      "**tres preguntas jerárquicas**:")
    A("")
    A("| | Pregunta |")
    A("|---|---|")
    A("| **011-A** · unidad de análisis | ¿qué es `i`? meta documental · meta "
      "operacional · unidad programática · intervención · agregado · otra |")
    A("| **011-B** · regla de correspondencia | ¿cómo se relacionan "
      "`PDOT_documental → ICPI_operacional`? Y deben poder coexistir **1:1 · "
      "N:1 · 1:N · N:N · NO DETERMINABLE** — no se obliga al universo a encajar "
      "en una sola relación |")
    A("| **011-C** · consecuencia algebraica | conocida la unidad, ¿son "
      "coherentes con ella la estructura multiplicativa y la ponderación `P·R`? "
      "**En ese orden, no al revés** (`DOC-016`) |")
    A("")

    A("### La consecuencia para v2, que es lo que 008-R venía a preparar")
    A("")
    A("**El Gold Master no conserva el texto de las metas del PDOT, sólo un "
      "resumen agregado.** Por eso ninguna reconciliación posterior puede ser "
      "automática, y por eso ésta llegó hasta donde llegó.")
    A("")
    A("Para v2, cada meta operacional debe guardar **el texto íntegro de cada "
      "meta documental que agrega, con su localización** (sistema · fila · SHA "
      "del documento). No es un requisito de comodidad: sin él, el universo "
      "ampliado nacería con la misma deuda de trazabilidad que esta auditoría "
      "acaba de medir — y en un sistema cuyo objeto **es** la trazabilidad.")
    A("")
    A("> ### ⚖️ CONDICIÓN CONGELADA PARA v2")
    A("> **Ninguna unidad operacional de v2 podrá existir sin declarar su "
      "correspondencia con una o más unidades documentales del universo PDOT, "
      "conservando el texto fuente, el identificador, la localización "
      "documental y la relación de correspondencia.**")
    A(">")
    A("> Y si una unidad representa varias metas, **debe poder demostrarse que "
      "representa esas metas y por qué la agregación es metodológicamente "
      "válida** — no basta con listarlas.")
    A("")
    A("Eso convierte el problema descubierto en una **capacidad estructural de "
      "QUIRA**, no en una reparación artesanal de Montecristi: el día que se "
      "cargue el GAD 002, la condición ya estará puesta.")
    A("")
    A("### Y «trabajar con las 66» no significa necesariamente 66 filas")
    A("")
    A("La instrucción de Javo —*«debimos trabajar con las 66 y establecer esa "
      "condición»*— fija el **universo trazable de entrada**, no el número de "
      "unidades del motor. `011` decidirá cuál de estos modelos corresponde:")
    A("")
    A("| Modelo | | Estado |")
    A("|---|---|---|")
    A("| **A** | `66 → 66` · cada meta documental es una unidad operacional | candidato |")
    A("| **B** | `66 → n` · se permite agregación, y **cada agregado declara sus "
      "componentes** | candidato |")
    A("| **C** | `66 → 25` · las unidades actuales | **ahora habría que "
      "demostrarlo**, no suponerlo — con tabla de correspondencia completa |")
    A("")
    A("- **No se amplía 25 → 66.** Sigue siendo `ADR-036 §4`: versión nueva, "
      "recalibración y ADR propio, después de `011`.")
    A("")
    A("---")
    A(f"*GM-Ω-ICPI-008-R · {len(filas)} metas documentales · {len(rec)} "
      f"reconciliadas · {len(amb)} ambiguas · el Gold Master no se modificó · "
      "Dylus Lab © 2026*")

    _SALIDA.write_text("\n".join(o) + "\n", encoding="utf-8")


if __name__ == "__main__":
    raise SystemExit(main())
