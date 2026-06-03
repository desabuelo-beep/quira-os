# -*- coding: utf-8 -*-
"""
chunker_holding.py — DOCX/PDF → chunks seccion-nivel
QUIRA Gov · Holding Municipal Montecristi · Dylus Lab © 2026

A diferencia del chunker normativo (articulo-nivel), este chunker opera
sobre documentos de gestion publica (RC, PP, POA, PAC, SIGAD) que tienen
estructura de secciones/capitulos, no de articulos juridicos.

Estrategia:
  1. Extraer texto del DOCX (python-docx) o PDF (pdfplumber)
  2. Detectar limites de seccion con regex flexible
  3. Chunks de ~400 palabras con overlap de 2 parrafos
  4. SHA256 por chunk para idempotencia

Dependencias: python-docx, pdfplumber
"""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

# ── REGEX CABECERAS DE SECCION ────────────────────────────────────────────────
# Detecta: "CAPITULO I", "SECCIÓN 3", "1.", "1.1", "I.", titulos en mayuscula
SECTION_RE = re.compile(
    r"(?:^|\n)\s*"
    r"(?:"
    r"CAP[IÍ]TULO\s+[\dIVXivx]+"            # CAPÍTULO I / CAPÍTULO 1
    r"|SECCI[OÓ]N\s+[\dIVXivx]+"             # SECCIÓN 3
    r"|\d+\.\d*\s+[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ ]+" # 1.2 TITULO EN MAYUSCULAS
    r"|\d+\.\s+[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ ]+"    # 1. TITULO EN MAYUSCULAS
    r"|[IVX]+\.\s+[A-ZÁÉÍÓÚÑ]"              # I. TITULO
    r"|[A-ZÁÉÍÓÚÑ]{4,}(?:\s+[A-ZÁÉÍÓÚÑ]{2,}){0,5}\s*$" # LINEA TODO MAYUSCULAS
    r")",
    re.MULTILINE
)

CHUNK_MAX_WORDS    = 400

# Overlap adaptativo por tipo de documento (fix OBS-004)
# DOCX: 10 palabras — los docs de gestion tienen secciones bien delimitadas.
#        50 palabras causaba que el header completo se repitiera en cada chunk.
# PDF:  50 palabras — texto continuo sin separacion natural de secciones.
CHUNK_OVERLAP_DOCX = 10
CHUNK_OVERLAP_PDF  = 50

# Umbral deduplicacion de prefijo comun entre chunks consecutivos (OBS-004)
DEDUPE_PREFIX_THRESHOLD = 15  # si >15 palabras iniciales iguales, se stripea


@dataclass
class HoldingChunk:
    """Un chunk de documento del Holding Municipal."""
    seccion_raw:   str               # cabecera de sección detectada o "PÁRRAFO"
    chunk_seq:     int               # secuencia en el documento
    contenido:     str               # texto del chunk
    palabras:      int               # conteo de palabras
    sha256:        str               # hash del contenido
    pagina_inicio: Optional[int] = None  # para PDFs

    @classmethod
    def from_text(cls, text: str, seq: int, seccion: str = "CONTENIDO",
                  pagina: Optional[int] = None) -> "HoldingChunk":
        text = text.strip()
        sha = hashlib.sha256(text.encode("utf-8")).hexdigest()
        return cls(
            seccion_raw   = seccion,
            chunk_seq     = seq,
            contenido     = text,
            palabras      = len(text.split()),
            sha256        = sha,
            pagina_inicio = pagina,
        )


# ── EXTRACCIÓN DE TEXTO ────────────────────────────────────────────────────────

def _extract_text_docx(path: Path) -> list[tuple[str, int]]:
    """
    Extrae párrafos de un DOCX. Si el documento no tiene párrafos con texto
    (ej. POA/PAC enteramente en tablas), extrae filas de tablas como fallback.
    Retorna: lista de (texto, numero_secuencia)
    """
    from docx import Document
    doc = Document(str(path))
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append((text, i))

    # Documentos con >= 5 párrafos tienen contenido narrativo real — usar párrafos.
    # Documentos con < 5 párrafos son planillas basadas en tablas — usar tabla-fallback.
    if len(paragraphs) >= 5:
        return paragraphs

    # Fallback: documento basado en tablas (POA/PAC tipo planilla)
    rows: list[tuple[str, int]] = []
    seq = 0
    for table in doc.tables:
        for row in table.rows:
            # python-docx repite celdas en celdas fusionadas — deduplicar
            seen: set[str] = set()
            unique: list[str] = []
            for cell in row.cells:
                ct = cell.text.strip()
                if ct and ct not in seen:
                    seen.add(ct)
                    unique.append(ct)
            if unique:
                rows.append((" | ".join(unique), seq))
                seq += 1
    return rows


def _extract_text_pdf(path: Path) -> list[tuple[str, int]]:
    """
    Extrae texto de un PDF página por página usando pdfplumber.
    Retorna: lista de (texto_pagina, numero_pagina)
    """
    import pdfplumber
    pages = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages.append((text, i))
    return pages


def _extract_text_txt(path: Path) -> list[tuple[str, int]]:
    """
    Extrae líneas de un .txt pre-procesado (ej. extraído de Drive via Chrome).
    Retorna: lista de (linea, numero_linea)
    """
    lines = []
    with open(str(path), encoding="utf-8") as f:
        for i, line in enumerate(f):
            text = line.strip()
            if text and len(text) > 4:
                lines.append((text, i))
    return lines


def _extract_text(path: Path) -> tuple[list[tuple[str, int]], str]:
    """
    Auto-detecta formato y extrae texto.
    Retorna: (fragmentos, tipo) donde tipo = 'docx' | 'pdf' | 'txt'
    """
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_text_docx(path), "docx"
    elif suffix == ".pdf":
        return _extract_text_pdf(path), "pdf"
    elif suffix == ".txt":
        return _extract_text_txt(path), "txt"
    else:
        raise ValueError(f"Formato no soportado: {suffix}")


# ── CHUNKING ──────────────────────────────────────────────────────────────────

def _dedupe_common_prefix(chunks: list[HoldingChunk]) -> list[HoldingChunk]:
    """
    Elimina prefijos comunes repetidos entre chunks consecutivos (fix OBS-004).

    Problema: En DOCX de gestión pública (RC, POA), el template acumula la
    jerarquía del documento en cada párrafo. Ej: cada chunk comienza con
    "INFORME RC N°17649 PERIODO 2023 > CUMPLIMIENTO > METAS > ..."
    sin que eso aporte valor semántico adicional.

    Solución: si dos chunks consecutivos comparten >= DEDUPE_PREFIX_THRESHOLD
    palabras iniciales, se stripea el prefijo del segundo chunk.
    """
    if len(chunks) < 2:
        return chunks

    result = [chunks[0]]
    for i in range(1, len(chunks)):
        prev_words = result[-1].contenido.split()
        curr_words = chunks[i].contenido.split()

        # Contar palabras comunes iniciales
        common = 0
        for j in range(min(len(prev_words), len(curr_words), 80)):
            if prev_words[j].lower() == curr_words[j].lower():
                common += 1
            else:
                break

        if common >= DEDUPE_PREFIX_THRESHOLD:
            stripped = " ".join(curr_words[common:]).strip()
            if len(stripped.split()) >= 15:  # mantener chunk solo si queda contenido util
                fixed = HoldingChunk.from_text(
                    stripped,
                    chunks[i].chunk_seq,
                    chunks[i].seccion_raw,
                    chunks[i].pagina_inicio,
                )
                result.append(fixed)
            # si queda muy corto, descartar el chunk (era puro header)
        else:
            result.append(chunks[i])

    return result


def _split_into_chunks(fragments: list[tuple[str, int]],
                       doc_type: str) -> list[HoldingChunk]:
    """
    Divide los fragmentos en chunks de ~CHUNK_MAX_WORDS palabras.
    Detecta cabeceras de sección como límites naturales.

    Overlap adaptativo (fix OBS-004):
      DOCX: 10 palabras — evita arrastre de cabeceras en templates de gestion.
      PDF:  50 palabras — texto continuo, el overlap preserva contexto.
    """
    # Seleccionar overlap segun tipo de documento
    chunk_overlap = CHUNK_OVERLAP_DOCX if doc_type == "docx" else CHUNK_OVERLAP_PDF

    chunks: list[HoldingChunk] = []
    current_words: list[str] = []
    current_seccion = "CONTENIDO"
    current_page: Optional[int] = None
    seq = 0

    def flush(words: list[str], seccion: str, page: Optional[int]) -> Optional[HoldingChunk]:
        text = " ".join(words).strip()
        if len(text) < 20:   # ignorar fragmentos muy cortos
            return None
        return HoldingChunk.from_text(text, seq, seccion, page)

    for fragment, page_or_para in fragments:
        page_ref = page_or_para if doc_type == "pdf" else None

        # Detectar si el fragmento es una cabecera de sección
        is_section_header = bool(SECTION_RE.match(fragment)) or (
            len(fragment) < 120 and fragment.isupper() and len(fragment.split()) >= 2
        )

        if is_section_header and current_words:
            # Emitir chunk actual antes de nueva sección
            chunk = flush(current_words, current_seccion, current_page)
            if chunk:
                chunks.append(chunk)
                seq += 1
            # Overlap reducido para DOCX: no arrastrar headers de sección
            overlap_words = current_words[-chunk_overlap:] if len(current_words) > chunk_overlap else current_words[:]
            current_words = overlap_words
            current_seccion = fragment[:100]  # nueva sección
            current_page = page_ref

        # Agregar texto al buffer
        fragment_words = fragment.split()
        current_words.extend(fragment_words)

        # Emitir si supera el límite
        while len(current_words) >= CHUNK_MAX_WORDS:
            chunk_words = current_words[:CHUNK_MAX_WORDS]
            chunk = flush(chunk_words, current_seccion, current_page)
            if chunk:
                chunks.append(chunk)
                seq += 1
            # Overlap adaptativo
            current_words = current_words[CHUNK_MAX_WORDS - chunk_overlap:]
            current_page = page_ref

    # Emitir lo que queda
    if current_words:
        chunk = flush(current_words, current_seccion, current_page)
        if chunk:
            chunks.append(chunk)

    # Post-proceso: deduplicar prefijos comunes (solo DOCX — PDFs ya no tienen este problema)
    if doc_type == "docx" and len(chunks) > 1:
        chunks = _dedupe_common_prefix(chunks)

    return chunks


# ── API PÚBLICA ───────────────────────────────────────────────────────────────

def chunk_holding_doc(path: str | Path) -> list[HoldingChunk]:
    """
    Chunker principal para documentos del Holding Municipal.
    Acepta DOCX o PDF. Retorna lista de HoldingChunk.

    Uso:
        chunks = chunk_holding_doc("path/to/RC-GAD-2024.docx")
        for c in chunks:
            print(c.seccion_raw, c.palabras, c.sha256[:8])
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Archivo no encontrado: {path}")

    fragments, doc_type = _extract_text(path)
    chunks = _split_into_chunks(fragments, doc_type)
    return chunks


# ── TEST ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Uso: python chunker_holding.py <path_to_docx_or_pdf>")
        sys.exit(1)

    target = Path(sys.argv[1])
    print(f"Procesando: {target.name}")
    chunks = chunk_holding_doc(target)
    print(f"Chunks generados: {len(chunks)}")
    for c in chunks[:3]:
        print(f"  [{c.chunk_seq}] {c.seccion_raw[:50]:50s} | {c.palabras:4d} palabras | {c.sha256[:12]}")
    if len(chunks) > 3:
        print(f"  ... {len(chunks)-3} chunks mas")
