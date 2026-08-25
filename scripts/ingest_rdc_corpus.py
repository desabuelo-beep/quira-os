# -*- coding: utf-8 -*-
"""
ingest_rdc_corpus.py — Ingesta de informes de Rendición de Cuentas al corpus
QUIRA Gov · Dylus Lab © 2026

Los informes RDC del GAD (2023-2025) son documentos AUTOREPORTADOS del sujeto
observado: prosa + tablas, NO articulados como una norma. Se ingestan al corpus
Supabase (normativa_corpus) como clase documental propia para alimentar la
recuperación semántica y QUIRA IA (responder sobre el texto completo del informe:
"¿qué dijo el GAD sobre el agua potable en la rendición 2024?").

Reutiliza el pipeline normativo (scripts/normativa/ingest.py): mismo modelo local
de embeddings (paraphrase-multilingual-MiniLM-L12-v2, sin API), mismo insert
idempotente por SHA256 (ON CONFLICT DO NOTHING). El modelo corre local: el texto
del informe NUNCA entra a Claude (token-eficiente).

MODO DOCUMENTO: los informes no tienen "Art. N", así que se IGNORA el regex de
artículos del chunker (evita cortes espurios en citas legales tipo "Art. 88 LOPC")
y el texto se ventana en chunks <=450 palabras con solapamiento vía _make_chunk.

Uso:
  python scripts/ingest_rdc_corpus.py --dry            # cuenta chunks + estado DB, NO inserta
  python scripts/ingest_rdc_corpus.py                  # ingesta real (3 años)
  python scripts/ingest_rdc_corpus.py --sigla RC-GAD-2025   # solo un año

Idempotente: re-ejecutar no duplica (SHA256 UNIQUE).
"""
from __future__ import annotations

import argparse
import json
import os
import sys
import time
from pathlib import Path

try:
    from config import DATOS_DIR as _DATOS
except Exception:                                        # noqa: BLE001
    import os as _os
    from pathlib import Path as _P
    _DATOS = _P(_os.environ.get("QUIRA_DATOS", "."))

# ── Reutilizar el pipeline normativo ─────────────────────────────────────────
# Importar ingest.py parchea st.secrets desde .streamlit/secrets.toml y deja
# get_connection + helpers listos (mismo patrón probado del corpus normativo).
_NORM = Path(__file__).resolve().parent / "normativa"
sys.path.insert(0, str(_NORM))
import ingest as ing                       # noqa: E402  (patchea secrets al importar)
from chunker import _make_chunk            # noqa: E402  (document-mode window)
from docx import Document                  # noqa: E402

# ── Fuente: informes oficiales CPCCS (docx) · Holding Montecristi ─────────────
DOCX_BASE = (
    str(_DATOS / "Holding_Municipal_Montecristi" / "Rendición de cuentas 2023-2025" / "GAD Montecristi")
)

# NOTA: los nombres de archivo traen el typo original "Monteristi" (tal cual en disco).
ENTRIES = [
    {"sigla": "RC-GAD-2023", "archivo": "GAD Monteristi Rendición de cuentas 2023.docx",
     "nombre": "Informe de Rendición de Cuentas — GAD Montecristi 2023"},
    {"sigla": "RC-GAD-2024", "archivo": "GAD Monteristi Rendición de cuentas 2024.docx",
     "nombre": "Informe de Rendición de Cuentas — GAD Montecristi 2024"},
    {"sigla": "RC-GAD-2025", "archivo": "GAD Monteristi Rendición de cuentas 2025.docx",
     "nombre": "Informe de Rendición de Cuentas — GAD Montecristi 2025"},
]

# Metadatos comunes de la clase "informe de rendición" (no es norma).
COMMON = {
    "jerarquia":      5,                   # informes: por debajo de toda norma/instrumento
    "milestone_qlep": "RDC",               # bucket propio (no milestone QLEP normativo)
    "tipo_documento": "informe_rendicion",
    "dominios":       ["Dom09"],           # Rendición de Cuentas (su dominio nativo)
}


def _extract_text(path: str) -> str:
    """Texto plano del DOCX: párrafos + celdas de tabla (informes son table-heavy)."""
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in parts:
                    parts.append(t)
    return "\n".join(parts)


def _rows_for(entry: dict) -> tuple[list[dict], int]:
    """Chunkea un informe en modo documento → filas listas para normativa_corpus."""
    path = os.path.join(DOCX_BASE, entry["archivo"])
    full = _extract_text(path)
    chunks = _make_chunk("INFORME", None, full)   # bypass ARTICLE_RE, ventana por palabras
    file_sha = ing._file_sha256(path)
    rows = []
    for c in chunks:
        rows.append({
            "norma_sigla":    entry["sigla"],
            "norma_nombre":   entry["nombre"],
            "jerarquia":      COMMON["jerarquia"],
            "milestone_qlep": COMMON["milestone_qlep"],
            "tipo_documento": COMMON["tipo_documento"],
            "articulo_num":   None,
            "articulo_raw":   "INFORME",
            "chunk_seq":      c.chunk_seq,
            "contenido":      c.contenido,
            "palabras":       c.palabras,
            "dominios_quira": json.dumps(COMMON["dominios"], ensure_ascii=False),
            "sha256":         c.sha256,
            "archivo_nombre": Path(path).name,
            "archivo_sha256": file_sha,
        })
    return rows, len(full.split())


def _existing_siglas(conn) -> dict:
    """Cuenta chunks por sigla RC-GAD ya en el corpus (para reporte honesto).

    Usa el mismo patrón simple que _load_existing_hashes (sin GROUP BY ni alias,
    que el wrapper de conexión no digiere) y agrega en Python.
    """
    from collections import Counter
    # OJO: no usar LIKE 'RC-GAD%' — el '%' lo interpreta el driver psycopg2 como
    # marcador de parámetro ("tuple index out of range"). Filtramos en Python
    # (robusto en Postgres y SQLite), reusando el patrón de _load_existing_hashes.
    try:
        c = conn.cursor()
        c.execute("SELECT norma_sigla FROM normativa_corpus")
        cnt: Counter = Counter()
        for row in c.fetchall():
            try:
                sig = row["norma_sigla"]
            except (TypeError, KeyError, IndexError):
                sig = row[0]
            if sig and str(sig).startswith("RC-GAD"):
                cnt[sig] += 1
        return dict(cnt)
    except Exception as exc:
        print(f"   (aviso: no se pudo consultar siglas existentes: {exc})")
        return {}


def main() -> None:
    ap = argparse.ArgumentParser(description="Ingesta de informes RDC al corpus")
    ap.add_argument("--dry", action="store_true", help="cuenta y reporta, no inserta")
    ap.add_argument("--sigla", default=None, help="filtrar por sigla (RC-GAD-2025)")
    args = ap.parse_args()

    entries = [e for e in ENTRIES if (not args.sigla or e["sigla"] == args.sigla)]
    print("=" * 68)
    print("  QUIRA — Ingesta de informes de Rendición de Cuentas al corpus")
    print(f"  Documentos: {len(entries)}  ·  modo documento (chunk ≤450 palabras)")
    print("=" * 68)

    # [1] Chunkear
    allrows: list[dict] = []
    for e in entries:
        path = os.path.join(DOCX_BASE, e["archivo"])
        if not os.path.exists(path):
            print(f"  [MISSING] {e['sigla']:14s} {e['archivo']}")
            continue
        rows, nwords = _rows_for(e)
        print(f"  {e['sigla']:14s} {len(rows):4d} chunks · {nwords:,} palabras")
        allrows.extend(rows)
    print(f"\n  Total chunks: {len(allrows)}")

    # [2] Estado DB (lectura, seguro en dry)
    conn = ing.get_connection()
    print(f"  Conexión: modo={conn.mode}")
    existing_hashes = ing._load_existing_hashes(conn)
    ya = _existing_siglas(conn)
    print(f"  Corpus actual: {len(existing_hashes)} chunks totales · RC-GAD ya presentes: {ya or '—'}")
    new_rows = [r for r in allrows if r["sha256"] not in existing_hashes]
    print(f"  Nuevos a insertar: {len(new_rows)}  ({len(allrows) - len(new_rows)} ya en DB)")

    if args.dry:
        if allrows:
            ex = allrows[0]["contenido"][:220].replace("\n", " ")
            print(f"\n  Ejemplo chunk[0] ({allrows[0]['norma_sigla']}): {ex}…")
        print("\n  DRY-RUN — no se escribió nada.")
        conn.close()
        return

    if not new_rows:
        print("\n  Nada nuevo que insertar.")
        conn.close()
        return
    conn.close()

    # [3] Modelo + embeddings + insert (patrón probado del corpus normativo)
    print("\n  Cargando modelo de embeddings…")
    model = ing._get_model()
    print(f"  {ing.EMBED_MODEL} listo ({ing.EMBED_DIM} dim)")

    t0 = time.time()
    B = ing.BATCH_SIZE
    conn_ins = ing.get_connection()
    ins = 0
    for i in range(0, len(new_rows), B):
        batch = new_rows[i:i + B]
        vecs = ing._embed_batch(model, [r["contenido"] for r in batch])
        for r, v in zip(batch, vecs):
            if ing._insert_chunk(conn_ins, r, v):
                ins += 1
        print(f"    …{min(i + B, len(new_rows))}/{len(new_rows)} chunks")
    conn_ins.close()

    print(f"\n  OK — {ins} chunks insertados en {time.time() - t0:.1f}s.")
    print("=" * 68)


if __name__ == "__main__":
    main()
