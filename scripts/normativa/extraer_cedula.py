# -*- coding: utf-8 -*-
"""
scripts/normativa/extraer_cedula.py — la cédula presupuestaria como evidencia financiera
════════════════════════════════════════════════════════════════════════════════════════
POR QUÉ EXISTE (2026-08-12). `V_eSIGEF` se estuvo derivando del POA, y el POA **no
acredita devengado**: acredita planificación. Estar en el POA es haber previsto;
tener devengado es haber ejecutado. Son cosas distintas y una no prueba la otra.

La fuente que sí lo acredita son las **cédulas presupuestarias de gastos** que el
GAD y sus entidades publican como Conjunto de Datos del **Numeral 6 de LOTAIP**,
mensualmente. Traen `Codificado`, `Comprometido`, `Devengado` y `Pagado` por
partida — que es exactamente el criterio de `V_eSIGEF`.

QUÉ HACE: normaliza esos archivos a un registro por (entidad · período · partida)
**conservando la procedencia de cada cifra** —archivo, hoja, fila—, para que
cualquier valor derivado pueda reconstruirse hasta el documento que lo sostiene.

QUÉ NO HACE:
  · No suma, no promedia, no consolida meses. Cada corte se guarda como está.
  · No corrige cifras inconsistentes. Las copia y las deja visibles.
  · No infiere la partida cuando el archivo no la trae.

Uso:  python scripts/normativa/extraer_cedula.py [--anio 2025] [--json salida.json]
Dylus Lab © 2026
"""
from __future__ import annotations

import argparse
import json
import re
import sys
import unicodedata
from pathlib import Path

try:
    from config import DATOS_DIR as _DATOS
except Exception:                                        # noqa: BLE001
    import os as _os
    from pathlib import Path as _P
    _DATOS = _P(_os.environ.get("QUIRA_DATOS", "."))

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

BASE = Path(str(_DATOS / "Holding_Municipal_Montecristi" / "Cedulas Presupuestarias 2023-2026"))

# ── Las cuatro entidades del holding. El nombre de carpeta varía por año
#    («BOmberos», «BOMBEROS 2026»), así que se reconoce por raíz, no por igualdad.
ENTIDADES = {
    "gad": "GAD Montecristi", "bomber": "Cuerpo de Bomberos",
    "aseo": "Empresa Pública de Aseo", "patronat": "Patronato",
}

MESES = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio",
         "agosto", "septiembre", "octubre", "noviembre", "diciembre"]

# ══════════════════════════════════════════════════════════════════════════════
# DOS INSTRUMENTOS DISTINTOS, NO UNO CON VARIANTES
#
# 2023-2024 · «Cédula Presupuestaria de Gastos» jerárquica del eSIGEF. Trae el
#             CÓDIGO ESTRUCTURADO COMPLETO (`01.01.01.A100.110.2024.510105…`),
#             que identifica programa · proyecto · actividad · ítem.
# 2025-2026 · «Conjunto de datos» del Numeral 6 LOTAIP, plano. Trae SÓLO EL ÍTEM
#             de 6 dígitos: el devengado viene agregado y se pierde a qué
#             actividad correspondió.
#
# El instrumento nuevo tiene MENOS estructura que el que reemplazó. Se registra
# aquí porque condiciona hasta dónde puede atribuirse la ejecución (ver §
# `cruce_poa_cedula.py`), y es el mismo patrón que OBS-027 halló en el POA.
# ══════════════════════════════════════════════════════════════════════════════
COLUMNAS = {
    # Numeral 6 (2025-2026) — las cuatro entidades usan el mismo encabezado;
    # sólo varían las tildes («Categoría»/«Categoria»), por eso se compara sin ellas.
    "cuenta": "partida", "categoria": "categoria", "descripcion": "descripcion",
    "asignado": "asignado", "modificado": "modificado", "codificado": "codificado",
    "monto certificado": "certificado", "comprometido": "comprometido",
    "devengado": "devengado", "pagado": "pagado",
    "porcentaje de ejecucion": "pct_ejecucion",
    # Cédula jerárquica (2023-2024)
    "codigo": "partida_completa", "estructura": "descripcion",
    "asignacion inicial": "asignado", "reformas": "modificado",
    "compromiso acumulado": "comprometido", "devengado acumulado": "devengado",
    "compromiso periodo": "comprometido_periodo", "devengado periodo": "devengado_periodo",
}
IMPORTES = {"asignado", "modificado", "codificado", "certificado", "comprometido",
            "devengado", "pagado", "pct_ejecucion", "comprometido_periodo",
            "devengado_periodo"}


def _sin_tildes(s: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFD", str(s))
                   if unicodedata.category(c) != "Mn").lower().strip()


def _num(v) -> float | None:
    """Importe, en cualquiera de las dos convenciones que conviven en la carpeta.

    El Numeral 6 escribe `1.866.275,79` (punto de miles, coma decimal) y la
    cédula del eSIGEF escribe `23,327,341.51` (al revés). **No se asume ninguna
    de las dos: manda el separador que aparece de último**, que es siempre el
    decimal. Fijar una convención habría devuelto `None` en silencio para todo
    un instrumento — un año entero de evidencia perdido sin un solo error.

    Devuelve `None` —no cero— cuando la celda está vacía o no es numérica. Un
    cero significa «se registró cero»; un vacío significa «no se registró».
    Confundirlos es exactamente el error que este módulo existe para evitar."""
    if v is None or (isinstance(v, str) and not v.strip()):
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip().replace("$", "").replace(" ", "").replace("%", "")
    if s in {"-", "--", "N/A", "n/a"}:
        return None
    ult_coma, ult_punto = s.rfind(","), s.rfind(".")
    if ult_coma > ult_punto:                      # 1.866.275,79
        s = s.replace(".", "").replace(",", ".")
    elif ult_punto > ult_coma:                    # 23,327,341.51
        s = s.replace(",", "")
    else:                                         # sin separadores
        s = s.replace(",", "").replace(".", "")
    try:
        return float(s)
    except ValueError:
        return None


def _partida(bruto) -> str:
    """El ítem de 6 dígitos, venga como venga escrito.

    Bomberos y Aseo lo publican punteado (`51.01.05`), el GAD corrido
    (`510105`) y la cédula jerárquica lo lleva dentro del código estructurado
    (`01.01.01.A100.110.2024.510105.000…`). Es el mismo clasificador."""
    s = str(bruto or "").strip()
    if not s:
        return ""
    if re.fullmatch(r"[\d.]+", s.replace(" ", "")) and s.count(".") <= 2:
        junto = s.replace(".", "").replace(" ", "")
        if re.fullmatch(r"\d{6}", junto):
            return junto
    for parte in s.split("."):                    # dentro del código estructurado
        if re.fullmatch(r"\d{6}", parte.strip()):
            return parte.strip()
    return ""


def _entidad(ruta: Path) -> str:
    """La entidad sale de la carpeta; si no, del propio nombre del archivo."""
    for parte in list(ruta.parts)[::-1]:
        n = _sin_tildes(parte)
        for raiz, nombre in ENTIDADES.items():
            if raiz in n:
                return nombre
    return "indeterminada"


def _periodo(ruta: Path) -> tuple[int | None, str | None]:
    n = _sin_tildes(ruta.name)
    anio = None
    for m in re.finditer(r"20(2[3-6])", ruta.name):
        anio = int(m.group(0))
        break
    if anio is None:                      # cae a la carpeta «Presupuestos 2024»
        for parte in ruta.parts[::-1]:
            m = re.search(r"20(2[3-6])", parte)
            if m:
                anio = int(m.group(0))
                break
    mes = next((m for m in MESES if m in n), None)
    if mes is None and "dic" in n:
        mes = "diciembre"
    return anio, mes


def _filas(ruta: Path) -> tuple[str, list[tuple]]:
    """Devuelve (nombre de hoja, filas) para .xlsx, .xls y .docx por igual.

    El `.docx` no es un descuido de lectura: **el Patronato publica su conjunto
    de datos del Numeral 6 dentro de un documento de Word** —el archivo se llama
    `…CONJUNTO DE DATOS.csv.docx`—. Es un dato abierto entregado en un formato
    que no lo es. Se lee la tabla igual, y el hecho queda registrado."""
    sufijo = ruta.suffix.lower()
    if sufijo == ".xls":
        import xlrd
        libro = xlrd.open_workbook(str(ruta))
        hoja = libro.sheet_by_index(0)
        return hoja.name, [tuple(hoja.row_values(i)) for i in range(hoja.nrows)]
    if sufijo == ".docx":
        import docx
        doc = docx.Document(str(ruta))
        if not doc.tables:
            return "sin tabla", []
        t = doc.tables[0]
        return "tabla Word 1", [tuple(c.text.strip() for c in fila.cells) for fila in t.rows]
    import openpyxl
    wb = openpyxl.load_workbook(ruta, data_only=True, read_only=True)
    ws = wb[wb.sheetnames[0]]
    filas = [tuple(f) for f in ws.iter_rows(values_only=True)]
    nombre = ws.title
    wb.close()
    return nombre, filas


def _cabecera(filas: list[tuple]) -> tuple[int, dict[int, str]] | None:
    """Localiza la fila de encabezado. No se asume que sea la primera: algunos
    archivos traen títulos, logotipos o filas en blanco arriba."""
    for i, f in enumerate(filas[:12]):
        mapa = {}
        for j, celda in enumerate(f):
            if celda is None:
                continue
            campo = COLUMNAS.get(_sin_tildes(celda))
            if campo:
                mapa[j] = campo
        vals = set(mapa.values())
        if vals & {"partida", "partida_completa"} and "devengado" in vals:
            return i, mapa
    return None


def extraer_archivo(ruta: Path) -> list[dict]:
    hoja, filas = _filas(ruta)
    cab = _cabecera(filas)
    if not cab:
        return []
    i_cab, mapa = cab
    anio, mes = _periodo(ruta)
    ent = _entidad(ruta)

    out: list[dict] = []
    for n, f in enumerate(filas[i_cab + 1:], i_cab + 2):   # nº de fila como lo ve Excel
        reg = {}
        for j, campo in mapa.items():
            if j >= len(f):
                continue
            reg[campo] = _num(f[j]) if campo in IMPORTES else (
                str(f[j]).strip() if f[j] is not None else "")
        # El ítem puede venir en su propia columna o dentro del código
        # estructurado. Lo que no resuelva a 6 dígitos es subtotal, encabezado
        # de grupo o fila de cierre: no es una partida y no se cuenta.
        partida = _partida(reg.get("partida") or reg.get("partida_completa"))
        if not partida:
            continue
        reg["partida"] = partida
        reg["_procedencia"] = {"archivo": ruta.name, "carpeta": ruta.parent.name,
                               "hoja": hoja, "fila": n,
                               "formato": ruta.suffix.lower().lstrip(".")}
        reg["entidad"], reg["anio"], reg["mes"] = ent, anio, mes
        out.append(reg)
    return out


def extraer_todo(anio: int | None = None) -> list[dict]:
    regs: list[dict] = []
    for ruta in sorted(BASE.rglob("*")):
        if ruta.suffix.lower() not in {".xlsx", ".xls", ".docx"} or ruta.name.startswith("~$"):
            continue
        if "poa" in _sin_tildes(ruta.name):        # POA de Bomberos: no es cédula
            continue
        a, _ = _periodo(ruta)
        if anio and a != anio:
            continue
        try:
            regs.extend(extraer_archivo(ruta))
        except Exception as e:                     # un archivo ilegible no detiene el resto
            print(f"  ⚠ no legible: {ruta.name} · {type(e).__name__}: {e}")
    return regs


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--anio", type=int, help="restringe a un año")
    ap.add_argument("--json", help="ruta donde volcar el resultado")
    args = ap.parse_args()

    print("CÉDULAS PRESUPUESTARIAS · Numeral 6 LOTAIP\n")
    regs = extraer_todo(args.anio)

    corte: dict[tuple, int] = {}
    for r in regs:
        corte[(r["entidad"], r["anio"], r["mes"])] = corte.get(
            (r["entidad"], r["anio"], r["mes"]), 0) + 1
    for (ent, a, mes), n in sorted(corte.items(), key=lambda x: (str(x[0][1]), x[0][0], str(x[0][2]))):
        print(f"  {str(a):5} {str(mes or '—'):11} {ent:26} {n:4} partidas")

    con_dev = sum(1 for r in regs if (r.get("devengado") or 0) > 0)
    print(f"\n  {len(regs)} registros · {len({r['partida'] for r in regs})} partidas distintas"
          f" · {con_dev} con devengado > 0")

    if args.json:
        Path(args.json).write_text(
            json.dumps(regs, ensure_ascii=False, indent=1), encoding="utf-8")
        print(f"  → {args.json}")


if __name__ == "__main__":
    main()
