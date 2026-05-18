"""
SENTINEL RAG · ocr_quality.py — Sprint 1.7
Scoring de calidad textual — detecta documentos mal escaneados o ilegibles.

No ejecuta OCR real — analiza la calidad del texto YA extraído.
Si el texto viene de un PDF escaneado con OCR deficiente, se detecta por:
  - Densidad baja de caracteres alfabéticos (mucho ruido)
  - Palabras demasiado largas (concatenación de tokens OCR)
  - Ausencia de marcadores del español institucional
  - Secuencias de artefactos OCR típicos (||, lll, @@)
  - Ratio texto/ruido

Soporta extracción desde: .txt, .md, .pdf (pypdf), .docx (python-docx), .xlsx (openpyxl)

Dylus Lab © 2026
"""
from __future__ import annotations
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

# ── UMBRALES ──────────────────────────────────────────────────────────────────
UMBRAL_OCR_ALTO   = 80    # >= 80  → calidad aceptable
UMBRAL_OCR_MEDIO  = 60    # 60-79  → advertencia
UMBRAL_CHAR_MIN   = 500   # G1: mínimo de caracteres

# Marcadores del español institucional ecuatoriano
SPANISH_MARKERS = [
    "el ", "la ", "los ", "las ", "de ", "del ", "en ", "con ",
    "por ", "para ", "que ", "una ", "son ", "como ", "este ",
    "según", "mediante", "considerando", "artículo", "presupuesto",
    "gobierno", "municipio", "cantón", "parroquia", "gestión",
    "planificación", "ejecución", "contratación", "ordenanza",
]

# Patrones de artefactos OCR típicos
OCR_ARTIFACT_PATTERNS = [
    r'\|{3,}',              # secuencias de pipe: |||
    r'_{6,}',               # subrayados largos: ______
    r'[l1I]{5,}',           # confusión l/1/I: llllll
    r'@{2,}',               # arrobas dobles: @@
    r'#{4,}',               # numerales: ####
    r'\x00+',               # bytes nulos
    r'([^\w\s,\.\-:;/])\1{3,}',  # cualquier carácter especial repetido 4+
    r'\b[A-Z]{15,}\b',      # palabras MAYÚSCULAS de 15+ chars (OCR)
    r'\b\w{25,}\b',         # palabras de 25+ chars (concatenación OCR)
]

# Fecha patterns (ES + ISO)
DATE_PATTERNS = [
    r'\b\d{1,2}/\d{1,2}/\d{4}\b',
    r'\b\d{1,2}-\d{1,2}-\d{4}\b',
    r'\b\d{4}-\d{2}-\d{2}\b',
    r'\b\d{1,2}\s+de\s+(?:enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre)\s+(?:del?\s+)?\d{4}\b',
    r'\b(?:enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre)\s+(?:del?\s+)?\d{4}\b',
]


# ── RESULTADO ─────────────────────────────────────────────────────────────────
@dataclass
class OcrQualityResult:
    score:              float   # 0-100 composite
    char_count:         int
    word_count:         int
    alpha_ratio:        float   # fracción de chars alfabéticos
    avg_word_length:    float   # longitud promedio de palabras
    spanish_detected:   bool    # marcadores del español encontrados
    artifacts_count:    int     # patrones OCR encontrados
    dates_found:        list[str]
    extraction_method:  str     # "txt", "md", "pdf", "docx", "xlsx", "direct"
    pasó_umbral:        bool    # score >= UMBRAL_OCR_ALTO
    nivel:              str     # "🟢 Alta" | "🟡 Media" | "🔴 Baja"
    detalles:           dict    = field(default_factory=dict)


# ── EXTRACCIÓN DE TEXTO ───────────────────────────────────────────────────────
def extract_text_from_file(filepath: Path) -> tuple[str, str]:
    """
    Extrae texto del archivo según su extensión.
    Retorna (texto, método_extracción).
    Las dependencias son opcionales — si no están instaladas, retorna vacío
    con el método indicando el problema.
    """
    ext = filepath.suffix.lower()

    if ext in (".txt", ".md"):
        try:
            return filepath.read_text(encoding="utf-8", errors="replace"), ext.lstrip(".")
        except Exception as e:
            return "", f"error-{ext}: {e}"

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(filepath))
            text_parts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
            return "\n".join(text_parts), "pdf"
        except ImportError:
            return "", "pdf-sin-pypdf"
        except Exception as e:
            return "", f"pdf-error: {e}"

    if ext in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(str(filepath))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            # También extraer tablas
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells if cell.text.strip()))
            return "\n".join(parts), "docx"
        except ImportError:
            return "", "docx-sin-python-docx"
        except Exception as e:
            return "", f"docx-error: {e}"

    if ext in (".xlsx", ".xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(filepath), read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(v) for v in row if v is not None)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts), "xlsx"
        except ImportError:
            return "", "xlsx-sin-openpyxl"
        except Exception as e:
            return "", f"xlsx-error: {e}"

    return "", f"formato-no-soportado({ext})"


# ── SCORING ───────────────────────────────────────────────────────────────────
def score_text_quality(text: str, filename: str = "") -> OcrQualityResult:
    """
    Calcula el score de calidad textual (0-100).

    Sistema de puntos:
      25 pts → alpha_ratio >= 0.65
      20 pts → avg_word_length entre 3 y 12 chars
      20 pts → español detectado
      20 pts → sin artefactos OCR
      15 pts → char_count >= 500
    """
    if not text or not text.strip():
        return OcrQualityResult(
            score=0, char_count=0, word_count=0, alpha_ratio=0,
            avg_word_length=0, spanish_detected=False, artifacts_count=0,
            dates_found=[], extraction_method="direct", pasó_umbral=False,
            nivel="🔴 Baja",
            detalles={"motivo": "Texto vacío"},
        )

    # Métricas básicas
    chars       = len(text)
    words       = text.split()
    alpha_chars = sum(1 for c in text if c.isalpha())
    alpha_ratio = alpha_chars / chars if chars > 0 else 0

    word_lengths   = [len(w) for w in words]
    avg_word_len   = sum(word_lengths) / len(word_lengths) if word_lengths else 0

    # Español detectado
    text_lower       = text.lower()
    spanish_hits     = sum(1 for m in SPANISH_MARKERS if m in text_lower)
    spanish_detected = spanish_hits >= 3

    # Artefactos OCR
    artifacts = 0
    for pattern in OCR_ARTIFACT_PATTERNS:
        artifacts += len(re.findall(pattern, text))

    # Fechas
    dates_found: list[str] = []
    for dp in DATE_PATTERNS:
        matches = re.findall(dp, text_lower)
        dates_found.extend(str(m) for m in matches if str(m) not in dates_found)
    dates_found = dates_found[:5]  # máximo 5

    # ── Score compuesto ───────────────────────────────────────────────────────
    score = 0.0

    # +25: ratio alfabético
    if alpha_ratio >= 0.65:
        score += 25
    elif alpha_ratio >= 0.50:
        score += 15
    elif alpha_ratio >= 0.35:
        score += 8

    # +20: longitud promedio de palabras (3-12 = natural; >15 = OCR concatenado)
    if 3.0 <= avg_word_len <= 12.0:
        score += 20
    elif 2.0 <= avg_word_len <= 15.0:
        score += 12
    elif avg_word_len > 0:
        score += 5

    # +20: español detectado
    if spanish_detected:
        score += 20
    elif spanish_hits >= 1:
        score += 10

    # +20: sin artefactos
    if artifacts == 0:
        score += 20
    elif artifacts <= 2:
        score += 12
    elif artifacts <= 5:
        score += 5

    # +15: longitud suficiente
    if chars >= UMBRAL_CHAR_MIN:
        score += 15
    elif chars >= 200:
        score += 8
    elif chars >= 50:
        score += 3

    score = round(min(score, 100.0), 1)
    pasó  = score >= UMBRAL_OCR_ALTO
    nivel = "🟢 Alta" if score >= UMBRAL_OCR_ALTO else ("🟡 Media" if score >= UMBRAL_OCR_MEDIO else "🔴 Baja")

    return OcrQualityResult(
        score            = score,
        char_count       = chars,
        word_count       = len(words),
        alpha_ratio      = round(alpha_ratio, 3),
        avg_word_length  = round(avg_word_len, 2),
        spanish_detected = spanish_detected,
        artifacts_count  = artifacts,
        dates_found      = dates_found,
        extraction_method = "direct",
        pasó_umbral      = pasó,
        nivel            = nivel,
        detalles         = {
            "spanish_hits":  spanish_hits,
            "long_words":    sum(1 for l in word_lengths if l > 20),
        },
    )


# ── PIPELINE COMPLETO ─────────────────────────────────────────────────────────
def score_file(filepath: str | Path) -> OcrQualityResult:
    """Extract + score in one call. Útil para tests y CLI."""
    p = Path(filepath)
    text, method = extract_text_from_file(p)
    result = score_text_quality(text, p.name)
    result.extraction_method = method
    return result


# ── CLI ───────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys, json

    if len(sys.argv) < 2:
        print("Uso: python -m sentinel.ocr_quality <archivo_o_texto>")
        sys.exit(1)

    target = sys.argv[1]
    path   = Path(target)

    if path.exists():
        r = score_file(path)
    else:
        # Tratar como texto directo
        r = score_text_quality(target)

    print(f"\n  Score OCR/Calidad:  {r.score:.0f}%  {r.nivel}")
    print(f"  Método extracción: {r.extraction_method}")
    print(f"  Caracteres:        {r.char_count}")
    print(f"  Palabras:          {r.word_count}")
    print(f"  Alpha ratio:       {r.alpha_ratio:.1%}")
    print(f"  Long. prom. word:  {r.avg_word_length:.1f} chars")
    print(f"  Español:           {'✅' if r.spanish_detected else '❌'}")
    print(f"  Artefactos OCR:    {r.artifacts_count}")
    print(f"  Fechas halladas:   {r.dates_found or 'ninguna'}")
    print(f"  Pasó umbral 80%:   {'✅' if r.pasó_umbral else '❌'}")
